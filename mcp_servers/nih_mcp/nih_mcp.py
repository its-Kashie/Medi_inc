# nih_mcp.py - NIH MCP Server (Report Tools Removed - Using External Report Generator)
"""
🏥 NIH (National Institute of Health) MCP Server - Updated
Handles: National Aggregation, WHO Proposals, 3-Year Analytics, R&D Emails
Reports: Delegated to report_generation_mcp_tool.py
"""

import os
import sys
import json
import time
from datetime import datetime
from functools import wraps
from email.message import EmailMessage
import smtplib
import threading
import asyncio

# ===== MCP Server =====
from mcp.server.fastmcp import FastMCP

os.environ['MCP_CLIENT_TIMEOUT'] = '10'


# ===== LAZY IMPORTS =====
def get_mongo():
    """Lazy MongoDB initialization"""
    global MongoClient, mongo_client, db
    if 'mongo_client' not in globals():
        try:
            from pymongo import MongoClient
            mongo_client = MongoClient("mongodb://localhost:27017/", serverSelectionTimeoutMS=3000)
            db = mongo_client["healthcare360_full"]
            mongo_client.admin.command('ping')
            debug_log("✅ MongoDB connected (NIH MCP)")
            return mongo_client, db
        except Exception as e:
            debug_log(f"⚠️ MongoDB unavailable: {e}")
            mongo_client = None
            db = None
            return None, None
    return mongo_client, db


def get_plotting():
    global pd, plt, matplotlib, io, base64
    if 'pd' not in globals():
        import pandas as pd
        import matplotlib
        matplotlib.use('Agg')
        import matplotlib.pyplot as plt
        import io
        import base64
    return pd, plt, matplotlib, io, base64


def get_email():
    global EmailMessage, smtplib, threading
    if 'EmailMessage' not in globals():
        from email.message import EmailMessage
        import smtplib
        import threading
    return EmailMessage, smtplib, threading


# ===== DEBUG LOGGING =====
def debug_log(msg):
    print(msg, file=sys.stderr, flush=True)


debug_log("✅ NIH MCP initialized (Reports delegated to external tool)")

# ===== TRACE LOG =====
TRACE_LOG = []


def log_trace(action: str, details: dict):
    entry = {
        "timestamp": datetime.now().isoformat(),
        "action": action,
        "details": details
    }
    TRACE_LOG.append(entry)
    debug_log(f"[NIH-MCP] {action}: {details}")


def traced_tool(func):
    @wraps(func)
    def wrapper(*args, **kwargs):
        call_id = f"{func.__name__}-{int(time.time() * 1000)}"
        log_trace("tool_start", {"tool": func.__name__, "call_id": call_id})
        try:
            result = func(*args, **kwargs)
            log_trace("tool_success", {"tool": func.__name__, "call_id": call_id})
            return result
        except Exception as e:
            log_trace("tool_error", {"tool": func.__name__, "error": str(e)})
            raise

    return wrapper


# ===== MCP SERVER =====
mcp = FastMCP("NIH Healthcare MCP - National Coordination")

# ===== DIRECTORIES =====
REPORTS_DIR = "generated_reports"
WHO_PROPOSALS_DIR = os.path.join(REPORTS_DIR, "who_proposals")
os.makedirs(REPORTS_DIR, exist_ok=True)
os.makedirs(WHO_PROPOSALS_DIR, exist_ok=True)

# ===== HOSPITALS & DEPARTMENTS =====
HOSPITALS_LIST = [
    "Mayo Hospital",
    "Services Hospital Lahore",
    "Jinnah Hospital Lahore",
    "Sir Ganga Ram Hospital",
    "Shalamar Hospital",
    "Lady Willingdon Hospital",
    "Fatima Memorial Hospital",
    "Shaukat Khanum Memorial",
    "PKLI Lahore",
    "Bahria International"
]

DEPARTMENTS = [
    "Cardiology",
    "Maternal Health",
    "Infectious Diseases",
    "Nutrition",
    "Mental Health",
    "NCD",
    "Endocrinology",
    "Oncology"
]

# ===== DEPARTMENT KEY MAPPING =====
DEPT_KEY_MAPPING = {
    "Cardiology": "cardiology_data",
    "Maternal Health": "maternal_health",
    "Infectious Diseases": "infectious_diseases",
    "Nutrition": "nutrition_data",
    "Mental Health": "mental_health_data",
    "NCD": "ncd_internal_medicine",
    "Endocrinology": "endocrinology_diabetes_data",
    "Oncology": "oncology_data"
}

# ===== EMAIL CONFIG =====
RESEARCH_SENDER_EMAIL = "nooreasal786@gmail.com"
APP_PASSWORD = "irph tole tuqr vfmi"


# ===== FOCAL PERSON FUNCTIONS =====
def get_university_focal_persons_from_excel() -> list:
    try:
        pd, _, _, _, _ = get_plotting()
        df = pd.read_excel("focal_persons_excels/university_focal_persons.xlsx")
        focal_persons = []
        for _, row in df.iterrows():
            focal_persons.append({
                "university": str(row['University']),
                "department": str(row['Department']),
                "name": str(row['Focal Person Name']),
                "email": str(row['Email']),
                "phone": str(row['Contact']),
                "notes": str(row.get('Notes / Internship Opportunities', ''))
            })
        return focal_persons
    except Exception as e:
        log_trace("university_focal_error", {"error": str(e)})
        return []


# ===== UTILITY FUNCTIONS =====
def get_quarter_dates(quarter: str, year: int) -> tuple:
    quarter_dates = {
        "Q1": (f"{year}-01-01", f"{year}-03-31"),
        "Q2": (f"{year}-04-01", f"{year}-06-30"),
        "Q3": (f"{year}-07-01", f"{year}-09-30"),
        "Q4": (f"{year}-10-01", f"{year}-12-31")
    }
    return quarter_dates.get(quarter, (f"{year}-01-01", f"{year}-03-31"))


# =============================================
# REPORT COORDINATION TOOLS (Using External MCP)
# =============================================

@mcp.tool()
@traced_tool
def request_department_report_generation(
        department: str,
        hospital: str,
        quarter: int,
        year: int
) -> dict:
    """
    Request a department report to be generated via the external Report Generation MCP.
    This tool coordinates report generation but doesn't generate reports itself.

    Args:
        department: Department name (e.g., 'Cardiology', 'Maternal Health')
        hospital: Hospital name
        quarter: Quarter number (1-4)
        year: Year

    Returns:
        dict with request status and metadata for agent to use
    """
    # Map department name to key
    dept_key = DEPT_KEY_MAPPING.get(department)

    if not dept_key:
        return {
            "status": "error",
            "message": f"Unknown department: {department}",
            "available_departments": list(DEPT_KEY_MAPPING.keys())
        }

    log_trace("report_request", {
        "department": department,
        "dept_key": dept_key,
        "hospital": hospital,
        "quarter": f"Q{quarter}",
        "year": year
    })

    # Return metadata for agent to call the external report generation MCP
    return {
        "status": "request_created",
        "message": f"Report request registered for {department} at {hospital}",
        "action_required": "Agent should call report_generation_mcp.generate_department_report_docx()",
        "parameters": {
            "department": dept_key,
            "hospital": hospital,
            "quarter": quarter,
            "year": year
        },
        "mcp_server": "report_generation_mcp",
        "tool_name": "generate_department_report_docx",
        "requested_at": datetime.now().isoformat()
    }


@mcp.tool()
@traced_tool
def request_all_departments_report_generation(
        hospital: str,
        quarter: int,
        year: int
) -> dict:
    """
    Request reports for ALL 8 departments for a hospital via external Report Generation MCP.

    Args:
        hospital: Hospital name
        quarter: Quarter number (1-4)
        year: Year

    Returns:
        dict with batch request metadata
    """
    log_trace("batch_report_request", {
        "hospital": hospital,
        "departments": len(DEPT_KEY_MAPPING),
        "quarter": f"Q{quarter}",
        "year": year
    })

    return {
        "status": "batch_request_created",
        "message": f"Batch report request registered for all 8 departments at {hospital}",
        "action_required": "Agent should call report_generation_mcp.generate_all_departments_for_hospital()",
        "parameters": {
            "hospital": hospital,
            "quarter": quarter,
            "year": year
        },
        "mcp_server": "report_generation_mcp",
        "tool_name": "generate_all_departments_for_hospital",
        "expected_reports": 8,
        "departments": list(DEPT_KEY_MAPPING.keys()),
        "requested_at": datetime.now().isoformat()
    }


# =============================================
# NATIONAL AGGREGATION TOOLS
# =============================================

@mcp.tool()
@traced_tool
def aggregate_national_statistics(department: str, quarter: str, year: int) -> dict:
    """
    Aggregate statistics for a department across all hospitals nationally.
    Uses MongoDB data directly (not Word reports).

    Args:
        department: Department name
        quarter: Quarter string (e.g., 'Q1')
        year: Year

    Returns:
        dict with national statistics
    """
    get_mongo()
    if mongo_client is None:
        return {"error": "MongoDB not available"}

    dept_key = DEPT_KEY_MAPPING.get(department)
    if not dept_key:
        return {"error": f"Unknown department: {department}"}

    # Collection mapping
    collection_map = {
        "cardiology_data": "cardiology_data",
        "maternal_health": "maternal_health",
        "infectious_diseases": "infectious_diseases",
        "nutrition_data": "nutrition_data",
        "mental_health_data": "mental_health_data",
        "ncd_internal_medicine": "ncd_internal_medicine",
        "endocrinology_diabetes_data": "endocrinology_diabetes_data",
        "oncology_data": "oncology_data"
    }

    collection_name = collection_map.get(dept_key)
    if not collection_name:
        return {"error": f"Collection not found for {dept_key}"}

    collection = db[collection_name]
    start_date, end_date = get_quarter_dates(quarter, year)

    # Determine date field
    date_field = "registration_date" if dept_key == "maternal_health" else "visit_date"

    # Aggregate across all hospitals
    hospital_stats = {}
    total_national = 0

    for hospital in HOSPITALS_LIST:
        query = {
            "hospital": hospital,
            date_field: {
                "$gte": datetime.fromisoformat(start_date),
                "$lte": datetime.fromisoformat(end_date)
            }
        }

        try:
            count = collection.count_documents(query)
            hospital_stats[hospital] = count
            total_national += count
        except Exception as e:
            hospital_stats[hospital] = 0
            log_trace("aggregation_error", {"hospital": hospital, "error": str(e)})

    return {
        "department": department,
        "period": f"{quarter}-{year}",
        "total_hospitals": len(HOSPITALS_LIST),
        "hospitals_reporting": len([h for h, c in hospital_stats.items() if c > 0]),
        "total_cases_national": total_national,
        "hospital_breakdown": hospital_stats,
        "average_per_hospital": round(total_national / len(HOSPITALS_LIST), 2),
        "generated_at": datetime.now().isoformat()
    }


@mcp.tool()
@traced_tool
def aggregate_all_departments_national(quarter: str, year: int) -> dict:
    """
    Aggregate statistics for ALL departments nationally.

    Args:
        quarter: Quarter string (e.g., 'Q1')
        year: Year

    Returns:
        dict with comprehensive national statistics
    """
    all_dept_stats = {}
    total_patients_all = 0

    for dept_name in DEPARTMENTS:
        stats = aggregate_national_statistics(dept_name, quarter, year)
        if "error" not in stats:
            all_dept_stats[dept_name] = stats
            total_patients_all += stats.get("total_cases_national", 0)

    return {
        "period": f"{quarter}-{year}",
        "total_departments": len(DEPARTMENTS),
        "departments_reporting": len(all_dept_stats),
        "total_patients_all_departments": total_patients_all,
        "department_breakdown": all_dept_stats,
        "generated_at": datetime.now().isoformat()
    }


# =============================================
# 3-YEAR TREND ANALYSIS
# =============================================

@mcp.tool()
@traced_tool
def analyze_three_year_trends(department: str, hospital: str) -> dict:
    """
    Analyze 3-year trends for a department at a specific hospital.

    Args:
        department: Department name
        hospital: Hospital name

    Returns:
        dict with 3-year trend data
    """
    get_mongo()
    if mongo_client is None:
        return {"error": "MongoDB not available"}

    dept_key = DEPT_KEY_MAPPING.get(department)
    if not dept_key:
        return {"error": f"Unknown department: {department}"}

    years = [2023, 2024, 2025]
    quarters = ["Q1", "Q2", "Q3", "Q4"]
    all_data = []

    collection_map = {
        "cardiology_data": "cardiology_data",
        "maternal_health": "maternal_health",
        "infectious_diseases": "infectious_diseases",
        "nutrition_data": "nutrition_data",
        "mental_health_data": "mental_health_data",
        "ncd_internal_medicine": "ncd_internal_medicine",
        "endocrinology_diabetes_data": "endocrinology_diabetes_data",
        "oncology_data": "oncology_data"
    }

    collection = db[collection_map[dept_key]]
    date_field = "registration_date" if dept_key == "maternal_health" else "visit_date"

    for year in years:
        for quarter in quarters:
            start_date, end_date = get_quarter_dates(quarter, year)
            query = {
                "hospital": hospital,
                date_field: {
                    "$gte": datetime.fromisoformat(start_date),
                    "$lte": datetime.fromisoformat(end_date)
                }
            }

            try:
                count = collection.count_documents(query)
                if count > 0:
                    all_data.append({
                        "period": f"{quarter}-{year}",
                        "year": year,
                        "quarter": quarter,
                        "total_patients": count
                    })
            except Exception as e:
                log_trace("three_year_analysis_error", {"period": f"{quarter}-{year}", "error": str(e)})

    total_patients_3yr = sum(d['total_patients'] for d in all_data)
    avg_per_quarter = total_patients_3yr / len(all_data) if all_data else 0

    return {
        "department": department,
        "hospital": hospital,
        "years_analyzed": "2023-2025",
        "total_quarters": len(all_data),
        "total_patients_3yr": total_patients_3yr,
        "avg_per_quarter": round(avg_per_quarter, 2),
        "quarterly_breakdown": all_data,
        "generated_at": datetime.now().isoformat()
    }


# =============================================
# WHO PROPOSAL GENERATOR
# =============================================

@mcp.tool()
@traced_tool
def generate_who_proposal(research_area: str, national_data: dict) -> dict:
    """
    Generate a WHO research proposal based on national data.

    Args:
        research_area: Research focus area
        national_data: National statistics dict

    Returns:
        dict with proposal details
    """
    proposal_id = f"WHO-PROP-{int(time.time())}"
    proposal = {
        "proposal_id": proposal_id,
        "title": f"Indigenous Health Research: {research_area} in Pakistan",
        "executive_summary": f"""
**Problem Statement:** National healthcare data reveals critical gaps in {research_area} across 10 major hospitals in Lahore.

**Evidence Base:** Analysis of {national_data.get('total_cases_national', 0)} patients over 3 years shows urgent need for intervention.
""",
        "disease_burden_analysis": {
            "data_source": "NIH Pakistan National Health Database",
            "hospitals_covered": 10,
            "patients_analyzed": national_data.get('total_cases_national', 0),
            "period": national_data.get('period', 'Unknown'),
            "key_findings": national_data.get('hospital_breakdown', {})
        },
        "funding_request": {
            "amount_usd": 500000,
            "duration_months": 24,
            "breakdown": {
                "research_staff": 200000,
                "equipment": 150000,
                "training": 100000,
                "overhead": 50000
            }
        },
        "intervention_plan": [
            f"Expand {research_area} screening programs",
            "Train 100+ healthcare workers",
            "Develop culturally-appropriate intervention models",
            "Establish monitoring systems"
        ],
        "expected_outcomes": [
            "Reduce mortality by 20% in target population",
            "Improve early detection rates by 30%",
            "Train 100+ healthcare workers",
            "Publish findings in international journals",
            "Create sustainable local solutions"
        ],
        "implementation_partners": [
            "National Institute of Health Pakistan",
            "King Edward Medical University",
            "University of Health Sciences Lahore",
            "Fatima Jinnah Medical University"
        ],
        "generated_at": datetime.now().isoformat(),
        "status": "draft_ready_for_review"
    }

    proposal_file = os.path.join(WHO_PROPOSALS_DIR, f"{proposal_id}.json")
    with open(proposal_file, 'w') as f:
        json.dump(proposal, f, indent=2, default=str)

    log_trace("who_proposal_generated", {"proposal_id": proposal_id, "research_area": research_area})

    return proposal


# =============================================
# R&D AGENT TOOLS
# =============================================
#
# @mcp.tool()
# @traced_tool
# def send_university_email_rnd(topic: str, content: str, attachments: list = None) -> dict:
#     """
#     Send research collaboration emails to university focal persons.
#
#     Args:
#         topic: Email subject topic
#         content: Email content
#         attachments: List of file paths to attach (optional)
#
#     Returns:
#         dict with email sending results
#     """
#     focal_persons = get_university_focal_persons_from_excel()
#     results = []
#     EmailMessage, smtplib, threading = get_email()
#
#     def send_email_thread(focal_person):
#         try:
#             msg = EmailMessage()
#             msg['Subject'] = f"Healthcare Research Collaboration: {topic}"
#             msg['From'] = RESEARCH_SENDER_EMAIL
#             msg['To'] = focal_person["email"]
#
#             email_body = f"""
# Dear {focal_person['name']},
#
# {content}
#
# **University:** {focal_person['university']}
# **Department:** {focal_person['department']}
#
# We invite collaboration on:
# - Internship opportunities
# - Final Year Projects (FYP)
# - Research partnerships
# - WHO-funded initiatives
#
# Best regards,
# National Institute of Health (NIH) Pakistan
# Research & Development Division
# """
#             msg.set_content(email_body)
#
#             if attachments:
#                 for attachment_path in attachments:
#                     if os.path.exists(attachment_path):
#                         with open(attachment_path, 'rb') as f:
#                             file_data = f.read()
#                             file_name = os.path.basename(attachment_path)
#                             msg.add_attachment(file_data, maintype='application', subtype='octet-stream',
#                                                filename=file_name)
#
#             with smtplib.SMTP("smtp.gmail.com", 587) as smtp:
#                 smtp.starttls()
#                 smtp.login(RESEARCH_SENDER_EMAIL, APP_PASSWORD)
#                 smtp.send_message(msg)
#
#             results.append({
#                 "university": focal_person["university"],
#                 "focal_person": focal_person["name"],
#                 "email": focal_person["email"],
#                 "status": "sent"
#             })
#         except Exception as e:
#             results.append({
#                 "university": focal_person["university"],
#                 "focal_person": focal_person["name"],
#                 "email": focal_person["email"],
#                 "status": "failed",
#                 "error": str(e)
#             })
#
#     threads = [threading.Thread(target=send_email_thread, args=(fp,)) for fp in focal_persons]
#     for t in threads:
#         t.start()
#     for t in threads:
#         t.join()
#
#     log_trace("university_emails_sent", {
#         "topic": topic,
#         "total": len(focal_persons),
#         "sent": len([r for r in results if r['status'] == 'sent'])
#     })
#
#     return {
#         "topic": topic,
#         "total_universities": len(focal_persons),
#         "emails_sent": len([r for r in results if r['status'] == 'sent']),
#         "emails_failed": len([r for r in results if r['status'] == 'failed']),
#         "details": results,
#         "sent_at": datetime.now().isoformat()
#     }
#
#
# @mcp.tool()
# @traced_tool
# def identify_research_priorities(national_data_all_depts: dict) -> dict:
#     """
#     Identify research priorities based on national data.
#
#     Args:
#         national_data_all_depts: Dict with all department national statistics
#
#     Returns:
#         dict with prioritized research areas
#     """
#     priorities = {
#         "high_priority": [],
#         "medium_priority": [],
#         "low_priority": [],
#         "recommendations": {
#             "urgent_interventions": [],
#             "research_needed": []
#         }
#     }
#
#     for dept, data in national_data_all_depts.items():
#         total_cases = data.get('total_cases_national', 0)
#         if total_cases > 5000:
#             priorities["high_priority"].append(f"{dept} - {total_cases} cases nationally")
#         elif total_cases > 2000:
#             priorities["medium_priority"].append(f"{dept} - {total_cases} cases")
#         else:
#             priorities["low_priority"].append(f"{dept} - On track")
#
#     priorities["recommendations"]["urgent_interventions"] = [
#         "Expand ICU capacity for cardiac emergencies",
#         "Launch maternal mortality reduction program",
#         "Water quality testing campaign"
#     ]
#
#     priorities["recommendations"]["research_needed"] = [
#         "Effectiveness of community health workers",
#         "Mental health screening models for Pakistan",
#         "Nutrition intervention impact studies"
#     ]
#
#     log_trace("research_priorities_identified", {
#         "high_priority_count": len(priorities["high_priority"])
#     })
#
#     return priorities


# =============================================
# UTILITY TOOLS
# =============================================

@mcp.tool()
@traced_tool
def list_nih_hospitals() -> dict:
    """List all hospitals in the system"""
    return {
        "total_hospitals": len(HOSPITALS_LIST),
        "hospitals": [
            {"id": f"H{i + 1:03d}", "name": hospital}
            for i, hospital in enumerate(HOSPITALS_LIST)
        ]
    }


@mcp.tool()
@traced_tool
def list_nih_departments() -> dict:
    """List all departments in the system"""
    return {
        "total_departments": len(DEPARTMENTS),
        "departments": [
            {"name": dept, "key": DEPT_KEY_MAPPING[dept]}
            for dept in DEPARTMENTS
        ]
    }


@mcp.tool()
@traced_tool
def receive_hospital_report(hospital_id: str, department: str, report_data: dict) -> dict:
    """
    Receive and log a hospital report submission.

    Args:
        hospital_id: Hospital identifier
        department: Department name
        report_data: Report data dict

    Returns:
        dict with receipt confirmation
    """
    log_trace("report_received", {
        "hospital_id": hospital_id,
        "department": department,
        "timestamp": datetime.now().isoformat()
    })

    return {
        "status": "received",
        "hospital_id": hospital_id,
        "department": department,
        "received_at": datetime.now().isoformat(),
        "report_id": f"RPT-{int(time.time())}"
    }


@mcp.tool()
@traced_tool
def validate_report_completeness(report_data: dict) -> dict:
    """
    Validate if a report has all required fields.

    Args:
        report_data: Report data dict

    Returns:
        dict with validation results
    """
    required_fields = ["hospital", "quarter", "year", "department", "total_patients", "statistics"]
    missing = [field for field in required_fields if field not in report_data]
    is_valid = len(missing) == 0

    return {
        "is_valid": is_valid,
        "missing_fields": missing,
        "completeness_score": ((len(required_fields) - len(missing)) / len(required_fields)) * 100
    }


@mcp.tool()
@traced_tool
def get_report_statistics(hospital: str, quarter: str, year: int) -> dict:
    """
    Get statistics for all departments at a hospital for a specific quarter.

    Args:
        hospital: Hospital name
        quarter: Quarter string (e.g., 'Q1')
        year: Year

    Returns:
        dict with statistics across all departments
    """
    get_mongo()
    if mongo_client is None:
        return {"error": "MongoDB not available"}

    stats = {}
    dept_collections = [
        ("cardiology_data", "visit_date"),
        ("maternal_health", "registration_date"),
        ("infectious_diseases", "visit_date"),
        ("nutrition_data", "visit_date"),
        ("mental_health_data", "visit_date"),
        ("ncd_internal_medicine", "visit_date"),
        ("endocrinology_diabetes_data", "visit_date"),
        ("oncology_data", "visit_date")
    ]

    start_date, end_date = get_quarter_dates(quarter, year)

    for dept_key, date_field in dept_collections:
        try:
            collection = db[dept_key]
            query = {
                "hospital": hospital,
                date_field: {
                    "$gte": datetime.fromisoformat(start_date),
                    "$lte": datetime.fromisoformat(end_date)
                }
            }
            count = collection.count_documents(query)
            stats[dept_key] = {"patient_count": count}
        except Exception as e:
            stats[dept_key] = {"patient_count": 0, "error": str(e)}

    return {
        "hospital": hospital,
        "quarter": quarter,
        "year": year,
        "department_statistics": stats,
        "total_patients": sum(s.get("patient_count", 0) for s in stats.values())
    }
#
# @mcp.tool()
# @traced_tool
# async def aggregate_hospital_departments_nih(hospital: str, quarter: str, year: int) -> dict:
#     """
#     Aggregate statistics for all departments at a specific hospital.
#     NOTE: This returns MongoDB statistics, NOT Word doc reports.
#
#     Args:
#         hospital: Hospital name
#         quarter: Quarter string (e.g., 'Q1')
#         year: Year
#
#     Returns:
#         dict with hospital-wide statistics
#     """
#     all_dept_stats = {}
#     total_patients = 0
#
#     for dept_name in DEPARTMENTS:
#         dept_key = DEPT_KEY_MAPPING.get(dept_name)
#         if not dept_key:
#             continue
#
#         get_mongo()
#         if mongo_client is None:
#             continue
#
#         collection = db[dept_key]
#         start_date, end_date = get_quarter_dates(quarter, year)
#         date_field = "registration_date" if dept_key == "maternal_health" else "visit_date"
#
#         query = {
#             "hospital": hospital,
#             date_field: {
#                 "$gte": datetime.fromisoformat(start_date),
#                 "$lte": datetime.fromisoformat(end_date)
#             }
#         }
#
#         try:
#             count = collection.count_documents(query)
#             all_dept_stats[dept_name] = {
#                 "patient_count": count,
#                 "department_key": dept_key
#             }
#             total_patients += count
#             await asyncio.sleep(0.5)  # Small delay
#         except Exception as e:
#             log_trace("dept_aggregation_error", {
#                 "department": dept_name,
#                 "hospital": hospital,
#                 "error": str(e)
#             })
#
#     return {
#         "hospital_id": f"H{hash(hospital) % 1000:03d}",
#         "hospital_name": hospital,
#         "quarter": quarter,
#         "year": year,
#         "departments": all_dept_stats,
#         "total_patients_all_depts": total_patients,
#         "departments_reporting": len([d for d in all_dept_stats.values() if d['patient_count'] > 0]),
#         "aggregated_at": datetime.now().isoformat()
#     }


# ===== WHO PROPOSAL GENERATOR (WORD DOCUMENT) =====
@mcp.tool()
@traced_tool
def generate_who_funding_proposal(
        research_area: str,
        national_data_summary: dict,
        three_year_trends: dict,
        output_filename: str = None
) -> dict:
    """
    Generate WHO funding proposal as downloadable Word document with graphs.

    Args:
        research_area: Research focus area (e.g., "Maternal Health Crisis")
        national_data_summary: Aggregated data from all 10 hospitals
        three_year_trends: 3-year trend analysis data
        output_filename: Custom filename (optional)

    Returns:
        dict with file path and proposal details
    """
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    import matplotlib.pyplot as plt
    import io

    # Create document
    doc = Document()

    # ===== TITLE PAGE =====
    title = doc.add_heading('WHO FUNDING PROPOSAL', level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    subtitle = doc.add_heading(f'Indigenous Health Research: {research_area}', level=1)
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph(f"\nSubmitted by: National Institute of Health (NIH) Pakistan")
    doc.add_paragraph(f"Date: {datetime.now().strftime('%B %d, %Y')}")
    doc.add_paragraph(f"Proposal ID: WHO-PROP-{int(time.time())}")



    # ===== EXECUTIVE SUMMARY =====
    doc.add_heading('Executive Summary', level=1)
    exec_para = doc.add_paragraph()
    exec_para.add_run("Problem Statement: ").bold = True
    exec_para.add_run(
        f"National healthcare data from 10 major hospitals reveals critical gaps in {research_area}. "
        f"Analysis of {national_data_summary.get('total_patients_all_departments', 0)} patients over 3 years "
        f"indicates urgent need for intervention."
    )

    doc.add_paragraph()
    exec_para2 = doc.add_paragraph()
    exec_para2.add_run("Evidence Base: ").bold = True
    exec_para2.add_run(
        f"Data collected from {national_data_summary.get('total_hospitals', 10)} hospitals across "
        f"{national_data_summary.get('total_departments', 8)} departments shows significant disease burden."
    )



    # ===== NATIONAL DATA ANALYSIS =====
    doc.add_heading('National Disease Burden Analysis', level=1)

    # Add summary table
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Light Grid Accent 1'

    headers = ['Department', 'Total Cases (National)', 'Hospitals Reporting']
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        cell.paragraphs[0].runs[0].font.bold = True

    # Add data rows
    for dept_name, dept_data in national_data_summary.get('department_breakdown', {}).items():
        row = table.add_row()
        row.cells[0].text = dept_name
        row.cells[1].text = str(dept_data.get('total_cases_national', 0))
        row.cells[2].text = str(dept_data.get('hospitals_reporting', 0))


    # ===== 3-YEAR TREND GRAPHS =====
    doc.add_heading('Three-Year Trend Analysis (2023-2025)', level=1)

    # Generate trend graph
    if three_year_trends and 'quarterly_breakdown' in three_year_trends:
        quarters = [d['period'] for d in three_year_trends['quarterly_breakdown']]
        patients = [d['total_patients'] for d in three_year_trends['quarterly_breakdown']]

        plt.figure(figsize=(10, 6))
        plt.plot(quarters, patients, marker='o', linewidth=2, markersize=8)
        plt.title(f"Patient Trends: {research_area} (2023-2025)", fontsize=14, fontweight='bold')
        plt.xlabel('Quarter', fontsize=12)
        plt.ylabel('Total Patients', fontsize=12)
        plt.xticks(rotation=45)
        plt.grid(True, alpha=0.3)
        plt.tight_layout()

        # Save graph to memory
        img_stream = io.BytesIO()
        plt.savefig(img_stream, format='png', dpi=300, bbox_inches='tight')
        img_stream.seek(0)
        plt.close()

        # Add graph to document
        doc.add_picture(img_stream, width=Inches(6))

        doc.add_paragraph()
        analysis = doc.add_paragraph()
        analysis.add_run("Key Finding: ").bold = True
        analysis.add_run(
            f"Total patients analyzed over 3 years: {three_year_trends.get('total_patients_3yr', 0)}. "
            f"Average per quarter: {three_year_trends.get('avg_per_quarter', 0):.0f} patients."
        )

    doc.add_page_break()

    # ===== FUNDING REQUEST =====
    doc.add_heading('Funding Request', level=1)

    funding_table = doc.add_table(rows=6, cols=2)
    funding_table.style = 'Light Grid Accent 1'

    funding_data = [
        ['Total Amount Requested', '$500,000 USD'],
        ['Project Duration', '24 months'],
        ['Research Staff', '$200,000'],
        ['Equipment & Infrastructure', '$150,000'],
        ['Training & Capacity Building', '$100,000'],
        ['Administrative Overhead', '$50,000']
    ]

    for i, (item, amount) in enumerate(funding_data):
        funding_table.rows[i].cells[0].text = item
        funding_table.rows[i].cells[1].text = amount
        if i == 0:
            funding_table.rows[i].cells[0].paragraphs[0].runs[0].font.bold = True
            funding_table.rows[i].cells[1].paragraphs[0].runs[0].font.bold = True


    # ===== INTERVENTION PLAN =====
    doc.add_heading('Proposed Interventions', level=1)

    interventions = [
        f"Expand {research_area} screening programs across all 10 hospitals",
        "Train 100+ healthcare workers in evidence-based protocols",
        "Develop culturally-appropriate intervention models for Pakistan",
        "Establish real-time monitoring and reporting systems",
        "Create community awareness campaigns"
    ]

    for intervention in interventions:
        p = doc.add_paragraph(intervention, style='List Bullet')
        p.paragraph_format.space_after = Pt(6)



    # ===== EXPECTED OUTCOMES =====
    doc.add_heading('Expected Outcomes', level=1)

    outcomes = [
        "Reduce mortality rates by 20% in target population",
        "Improve early detection rates by 30%",
        "Train 100+ healthcare professionals",
        "Publish findings in 3+ international peer-reviewed journals",
        "Create sustainable, locally-owned solutions",
        "Establish national disease registry"
    ]

    for outcome in outcomes:
        p = doc.add_paragraph(outcome, style='List Bullet')
        p.paragraph_format.space_after = Pt(6)



    # ===== IMPLEMENTATION PARTNERS =====
    doc.add_heading('Implementation Partners', level=1)

    partners = [
        "National Institute of Health (NIH) Pakistan - Lead Agency",
        "10 Major Teaching Hospitals in Lahore",
        "King Edward Medical University",
        "University of Health Sciences Lahore",
        "Fatima Jinnah Medical University",
        "Pakistan Medical Research Council"
    ]

    for partner in partners:
        doc.add_paragraph(partner, style='List Bullet')

    doc.add_page_break()

    # ===== MONITORING & EVALUATION =====
    doc.add_heading('Monitoring & Evaluation Plan', level=1)

    doc.add_paragraph("Quarterly Progress Reports:", style='List Bullet')
    doc.add_paragraph("Real-time data dashboard", style='List Bullet 2')
    doc.add_paragraph("Hospital-wise performance tracking", style='List Bullet 2')

    doc.add_paragraph("Annual Impact Assessment:", style='List Bullet')
    doc.add_paragraph("Mortality and morbidity analysis", style='List Bullet 2')
    doc.add_paragraph("Cost-effectiveness evaluation", style='List Bullet 2')

    doc.add_paragraph("Final Evaluation (Month 24):", style='List Bullet')
    doc.add_paragraph("Comprehensive impact study", style='List Bullet 2')
    doc.add_paragraph("Policy recommendations", style='List Bullet 2')

    doc.add_page_break()

    # ===== SUSTAINABILITY =====
    doc.add_heading('Sustainability Plan', level=1)

    doc.add_paragraph(
        "This project will establish permanent infrastructure and trained personnel "
        "within Pakistan's healthcare system. Government buy-in and local ownership "
        "ensure long-term sustainability beyond WHO funding period."
    )

    # ===== SIGNATURES =====

    doc.add_heading('Approval & Signatures', level=1)

    doc.add_paragraph(f"\n\n_______________________________")
    doc.add_paragraph("Director General, NIH Pakistan")
    doc.add_paragraph(f"Date: {datetime.now().strftime('%B %d, %Y')}")

    doc.add_paragraph(f"\n\n_______________________________")
    doc.add_paragraph("Secretary, Ministry of Health")
    doc.add_paragraph(f"Date: {datetime.now().strftime('%B %d, %Y')}")

    # ===== SAVE DOCUMENT =====
    if output_filename is None:
        safe_research_area = research_area.replace(" ", "_").replace("/", "-")
        output_filename = f"WHO_Proposal_{safe_research_area}_{datetime.now().strftime('%Y%m%d')}.docx"

    output_path = os.path.join(WHO_PROPOSALS_DIR, output_filename)
    doc.save(output_path)

    log_trace("who_proposal_generated", {
        "research_area": research_area,
        "file_path": output_path,
        "total_pages": len(doc.sections),
        "file_size_mb": round(os.path.getsize(output_path) / (1024 * 1024), 2)
    })

    return {
        "status": "success",
        "message": "WHO funding proposal generated successfully",
        "file_path": output_path,
        "filename": output_filename,
        "download_url": f"/api/proposals/download?file={output_filename}",
        "research_area": research_area,
        "total_hospitals": national_data_summary.get('total_hospitals', 10),
        "total_patients_analyzed": national_data_summary.get('total_patients_all_departments', 0),
        "funding_requested_usd": 500000,
        "generated_at": datetime.now().isoformat(),
        "file_size_mb": round(os.path.getsize(output_path) / (1024 * 1024), 2)
    }


if __name__ == "__main__":
    debug_log("🚀 NIH MCP Server Starting...")
    debug_log(f"✅ {len(HOSPITALS_LIST)} hospitals configured")
    debug_log(f"✅ {len(DEPARTMENTS)} departments configured")
    debug_log("📝 Reports delegated to: report_generation_mcp_tool.py")
    mcp.run()

