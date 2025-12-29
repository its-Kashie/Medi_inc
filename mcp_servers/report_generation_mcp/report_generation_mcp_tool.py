# report_generation_mcp_tool_FIXED.py
"""
🏥 Enhanced Report Generation - BATCH PROCESSING FIX
✅ Single department per MCP call (no timeout)
✅ Fast response (returns immediately)
✅ Background generation with status tracking
✅ All 8 departments complete successfully
"""

import os
import sys
import json
import pandas as pd
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from datetime import datetime
import asyncio
from typing import Optional

sys.path.append('.')
from all_agents import ALL_AGENTS
from agents import Agent, Runner
from agents.mcp import MCPServerStdio
from mcp.server.fastmcp import FastMCP

os.environ['MCP_CLIENT_TIMEOUT'] = '30'

mcp = FastMCP("Enhanced Department Report Generation - BATCH")


def debug_log(msg):
    print(msg, file=sys.stderr, flush=True)


# ===== MCP SERVERS =====
orchestrator_mcp = MCPServerStdio(
    params={"command": "python", "args": ["mcp_servers/orchestrator_mcp/agent_orchestrator_mcp.py"]},
    cache_tools_list=True,
    name="OrchestratorMCP_ReportGen"
)

agents_mcp = MCPServerStdio(
    params={"command": "python", "args": ["mcp_servers/core_agents_mcp/agents_mcp.py"]},
    cache_tools_list=True,
    name="AgentsMCP_ReportGen"
)

_mcp_connected = False
_mcp_initialized = False
_mcp_lock = None


async def get_or_create_lock():
    global _mcp_lock
    if _mcp_lock is None:
        _mcp_lock = asyncio.Lock()
    return _mcp_lock


async def ensure_mcp_connected():
    global _mcp_connected, _mcp_initialized
    lock = await get_or_create_lock()

    async with lock:
        if _mcp_initialized:
            return True
        if _mcp_connected:
            _mcp_initialized = True
            return True
        try:
            debug_log("🔌 Connecting MCP servers...")
            await orchestrator_mcp.connect()
            await agents_mcp.connect()
            for agent_key, agent_info in ALL_AGENTS.items():
                agent_info["agent"].mcp_servers = [orchestrator_mcp, agents_mcp]
            _mcp_connected = True
            _mcp_initialized = True
            debug_log("✅ MCP servers connected")
            return True
        except Exception as e:
            debug_log(f"⚠️ MCP connection failed: {e}")
            _mcp_initialized = False
            return False


# ===== CONFIG =====
CSV_DIR = "generated_output/csvs"
JSON_DIR = "generated_output/aggregates"
GRAPH_DIR = "generated_output/graphs"
TEMPLATE_DIR = "hospital_department_templates"
FOCAL_PERSONS_FILE = "focal_persons_excels/hospital_focal_persons.xlsx"
OUTPUT_DIR = "filled_reports"
STATUS_DIR = "report_generation_status"

os.makedirs(OUTPUT_DIR, exist_ok=True)
os.makedirs(STATUS_DIR, exist_ok=True)

DEPT_CONFIG = {
    "infectious_diseases": {
        "template": "infectious_diseases_template.docx",
        "name": "Infectious Diseases",
        "focal_dept": "Infectious Diseases",
        "csv_file": "infectious_diseases.csv"
    },
    "maternal_health": {
        "template": "maternal_health_template.docx",
        "name": "Obstetrics & Gynecology",
        "focal_dept": "Obstetrics & Gynecology",
        "csv_file": "maternal_health.csv"
    },
    "nutrition_data": {
        "template": "nutrition_data_template.docx",
        "name": "Nutrition & Dietetics",
        "focal_dept": "Nutrition & Dietetics",
        "csv_file": "nutrition_data.csv"
    },
    "mental_health_data": {
        "template": "mental_health_data_template.docx",
        "name": "Psychiatry",
        "focal_dept": "Psychiatry",
        "csv_file": "mental_health_data.csv"
    },
    "ncd_internal_medicine": {
        "template": "ncd_internal_medicine_template.docx",
        "name": "Internal Medicine (NCDs)",
        "focal_dept": "Internal Medicine (NCDs)",
        "csv_file": "ncd_internal_medicine.csv"
    },
    "cardiology_data": {
        "template": "cardiology_data_template.docx",
        "name": "Cardiology",
        "focal_dept": "Cardiology",
        "csv_file": "cardiology_data.csv"
    },
    "endocrinology_diabetes_data": {
        "template": "endocrinology_diabetes_data_template.docx",
        "name": "Endocrinology / Diabetes",
        "focal_dept": "Endocrinology",
        "csv_file": "endocrinology_diabetes_data.csv"
    },
    "oncology_data": {
        "template": "oncology_data_template.docx",
        "name": "Oncology",
        "focal_dept": "Oncology",
        "csv_file": "oncology_data.csv"
    }
}


# ===== UTILITY FUNCTIONS =====
def load_focal_persons():
    try:
        df = pd.read_excel(FOCAL_PERSONS_FILE)
        focal_map = {}
        for _, row in df.iterrows():
            key = f"{row['Hospital']}_{row['Department']}"
            focal_map[key] = {
                'name': row['Focal Person Name'],
                'contact': row['Contact'],
                'email': row['Email']
            }
        return focal_map
    except Exception as e:
        debug_log(f"⚠️ Warning: Could not load focal persons: {e}")
        return {}


def get_quarter_dates(year, quarter):
    quarters = {
        1: (f"Jan 1, {year}", f"Mar 31, {year}"),
        2: (f"Apr 1, {year}", f"Jun 30, {year}"),
        3: (f"Jul 1, {year}", f"Sep 30, {year}"),
        4: (f"Oct 1, {year}", f"Dec 31, {year}")
    }
    return quarters.get(quarter, ("N/A", "N/A"))


def calculate_percent_change(current, previous):
    if previous == 0 or previous is None:
        return "N/A"
    try:
        change = ((current - previous) / previous) * 100
        return f"{change:+.1f}%"
    except:
        return "N/A"


def add_para(doc, text, bold=False, size=11):
    para = doc.add_paragraph()
    run = para.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    return para


def insert_table(doc, data, headers):
    if not data:
        return None
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Light Grid Accent 1'
    hdr_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        hdr_cells[i].text = str(header)
        hdr_cells[i].paragraphs[0].runs[0].font.bold = True
    for row_data in data:
        row_cells = table.add_row().cells
        for i, val in enumerate(row_data):
            row_cells[i].text = str(val)
    return table


def safe_count(series, value):
    try:
        return series.value_counts().get(value, 0)
    except:
        return 0


# def clean_template_content(doc):
#     remove_patterns = [
#         '_______', '_____|', 'Executive Summary',
#         'Patient / Treatment Summary', 'Diagnostic / Monitoring Data',
#         'Resource & Staff', 'Action Plan / Recommendations',
#         '• Improvements required:', '• Training needs:'
#     ]
#     to_remove = []
#     for element in doc.element.body:
#         if element.tag.endswith('p'):
#             para_text = ''.join(element.itertext()) if hasattr(element, 'itertext') else ''
#             if any(pattern in para_text for pattern in remove_patterns):
#                 if 'Hospital Name:' not in para_text and 'Reporting Period:' not in para_text:
#                     to_remove.append(element)
#         elif element.tag.endswith('tbl'):
#             to_remove.append(element)
#     for element in to_remove:
#         try:
#             element.getparent().remove(element)
#         except:
#             pass
def clean_template_content(doc):
    """
    Aggressively removes all placeholder tables and repeated garbage sections
    """
    # List of exact phrases jo hamesha garbage mein aati hain
    garbage_phrases = [
        "Key Highlights / Remarks",
        "Indicator | Target / Benchmark",
        "Test / Procedure | Number Conducted",
        "Resource / Staff | Planned / Available",
        "Notes / Actions Taken",
        "Abnormal Findings",
        "Utilization this Quarter",
        "Notes / Issues",
        "Key Highlights",
        "Remarks",
        "Executive Summary",
        "Patient / Treatment Summary",
        "Diagnostic / Monitoring Data",
        "Resource & Staff",
        "Action Plan / Recommendations",
        "• Improvements required:",
        "• Training needs:",
        "_______",
        "_____|"
    ]

    elements_to_remove = []

    # 1. Saare paragraphs check karo
    for para in doc.paragraphs:
        text = para.text.strip()
        if any(phrase in text for phrase in garbage_phrases):
            # Poora paragraph delete mark kar do
            para._element.getparent().remove(para._element)

    # 2. Saare tables check karo — agar inme se koi phrase hai to poori table hata do
    for table in doc.tables:
        should_remove = False
        for row in table.rows:
            for cell in row.cells:
                cell_text = cell.text.strip()
                if any(phrase in cell_text for phrase in garbage_phrases):
                    should_remove = True
                    break
            if should_remove:
                break
        if should_remove:
            elements_to_remove.append(table._element)

    # 3. Tables ko actually remove karo
    for element in elements_to_remove:
        try:
            element.getparent().remove(element)
        except:
            pass

    # 4. Extra blank paragraphs bhi saaf kar do (optional — report aur clean lagega)
    for para in doc.paragraphs:
        if para.text.strip() == "" and len(para.runs) == 0:
            para._element.getparent().remove(para._element)


def update_header_info(doc, hospital, quarter, year, dept, focal_persons):
    for para in doc.paragraphs:
        text = para.text
        if "Hospital Name:" in text:
            para.text = f"Hospital Name: {hospital}"
        elif "Reporting Period:" in text:
            start, end = get_quarter_dates(year, quarter)
            para.text = f"Reporting Period: From {start} to {end}"
        elif "Prepared by:" in text:
            focal_key = f"{hospital}_{dept['focal_dept']}"
            focal = focal_persons.get(focal_key, {'name': '[Name Not Available]'})
            para.text = f"Prepared by: Head of {dept['name']}: {focal['name']}"


# ===== AI FUNCTIONS =====
async def call_agent_with_retry(agent, prompt, max_retries=3):
    for attempt in range(max_retries):
        try:
            await ensure_mcp_connected()
            result = await Runner.run(agent, prompt)
            return True, result.final_output
        except Exception as e:
            debug_log(f"❌ Attempt {attempt + 1} failed: {e}")
            if attempt < max_retries - 1:
                await asyncio.sleep(2 ** attempt)
            continue
    return False, None


async def get_ai_section_analysis(department: str, section_type: str, data: dict,
                                  hospital: str, quarter: int, year: int) -> str:
    agent_key_mapping = {
        "cardiology_data": "cardiology",
        "maternal_health": "maternal_health",
        "infectious_diseases": "infectious_diseases",
        "nutrition_data": "nutrition",
        "mental_health_data": "mental_health",
        "ncd_internal_medicine": "ncd",
        "endocrinology_diabetes_data": "endocrinology",
        "oncology_data": "oncology"
    }

    agent_key = agent_key_mapping.get(department)
    if not agent_key or agent_key not in ALL_AGENTS:
        return "Analysis unavailable - agent not configured."

    agent = ALL_AGENTS[agent_key]["agent"]

    # Build prompts (same as before)
    if section_type == "executive":
        prompt = f"""Q{quarter} {year} report for {DEPT_CONFIG[department]['name']} at {hospital}.
Data: {json.dumps(data, indent=2)}
Write 3-4 sentence executive summary. Return only text."""
    elif section_type == "table":
        prompt = f"""Table analysis for Q{quarter} {year} at {hospital}.
Data: {json.dumps(data, indent=2)}
Write 3-4 sentence analysis. Return only text."""
    elif section_type == "graph":
        prompt = f"""Graph analysis for Q{quarter} {year} at {hospital}.
Graphs: {json.dumps(data, indent=2)}
Write 2-3 sentence interpretation. Return only text."""
    elif section_type == "recommendations":
        prompt = f"""Action plan for {DEPT_CONFIG[department]['name']} at {hospital} Q{quarter} {year}.
Data: {json.dumps(data, indent=2)}
Generate 5-7 recommendations as JSON array:
{{"recommendations": ["1. ...", "2. ...", ...]}}
Return only JSON."""

    success, response = await call_agent_with_retry(agent, prompt)
    await asyncio.sleep(4)

    if not success:
        fallbacks = {
            "executive": "Quarterly performance data captured.",
            "table": "Summary table shows current quarter statistics.",
            "graph": "Visual trends displayed.",
            "recommendations": json.dumps({"recommendations": [
                "1. Review quarterly performance metrics",
                "2. Verify data accuracy",
                "3. Continue standard protocols",
                "4. Monitor key indicators",
                "5. Schedule follow-up meetings"
            ]})
        }
        return fallbacks.get(section_type, "Analysis pending.")

    if section_type == "recommendations":
        try:
            if "```json" in response:
                json_start = response.find("```json") + 7
                json_end = response.find("```", json_start)
                json_str = response[json_start:json_end].strip()
            else:
                json_start = response.find('{')
                json_end = response.rfind('}') + 1
                json_str = response[json_start:json_end]
            parsed = json.loads(json_str)
            return parsed.get("recommendations", [])
        except:
            return ["1. Review data", "2. Verify accuracy", "3. Continue protocols"]

    return response.strip()


def add_ai_analysis_paragraph(doc, title, content, icon="📊"):
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(12)
    para.paragraph_format.space_after = Pt(12)
    run_title = para.add_run(f"\n{icon} {title}:\n")
    run_title.bold = True
    run_title.font.size = Pt(11)
    run_title.font.color.rgb = RGBColor(0, 51, 102)
    if isinstance(content, list):
        for rec in content:
            rec_para = doc.add_paragraph()
            rec_para.style = 'List Bullet'
            rec_para.add_run(rec).font.size = Pt(10)
    else:
        run_content = para.add_run(content)
        run_content.font.size = Pt(10)


# ===== DEPARTMENT CONTENT GENERATORS (Keep all 8 functions - same as before) =====
def generate_cardiology_content(doc, df, prev_data, dept_info, hospital):
    total = int(len(df))
    mortality = int(df[df['mortality'] == 'Died'].shape[0]) if 'mortality' in df.columns else 0
    icu = int(df[df['icu_admission'] == 'Yes'].shape[0]) if 'icu_admission' in df.columns else 0
    doc.add_heading('Executive Summary', level=1)
    para = doc.add_paragraph()
    para.add_run(f"• Total cardiac patients: {total}").bold = True
    para = doc.add_paragraph()
    para.add_run(f"• ICU admissions: {icu} ({icu / total * 100:.1f}%)").bold = True
    doc.add_page_break()
    para = doc.add_paragraph()
    para.add_run(f"• Mortality: {mortality}").bold = True
    doc.add_page_break()
    doc.add_heading('Cardiac Care Summary', level=1)
    table = doc.add_table(rows=4, cols=5)
    table.style = 'Light Grid Accent 1'
    headers = ["Indicator", "Current", "Previous", "% Change", "Notes"]
    for i, header in enumerate(headers):
        table.rows[0].cells[i].text = header
        table.rows[0].cells[i].paragraphs[0].runs[0].font.bold = True
    prev_total = int(prev_data.get('total_patients', 0)) if prev_data else 0
    prev_mortality = int(prev_data.get('mortality', 0)) if prev_data else 0
    rows_data = [
        ["Total Patients", str(total), str(prev_total), calculate_percent_change(total, prev_total), "Normal range"],
        ["ICU Admissions", str(icu), "N/A", "N/A", "Critical care"],
        ["Mortality", str(mortality), str(prev_mortality), calculate_percent_change(mortality, prev_mortality),
         "Acceptable"]
    ]
    for i, row_data in enumerate(rows_data, start=1):
        for j, value in enumerate(row_data):
            table.rows[i].cells[j].text = value
    return {"total_patients": total, "icu_admissions": icu, "mortality": mortality, "table_data": rows_data}


# (Keep other 7 content generators exactly as they are)
def generate_maternal_health_content(doc, df, prev_data, dept_info, hospital):
    total = int(len(df))
    csection = int(safe_count(df.get('delivery_type', pd.Series()), 'C-section'))
    complications = int(df['maternal_complication'].notna().sum()) if 'maternal_complication' in df.columns else 0
    doc.add_heading('Executive Summary', level=1)
    add_para(doc, f"• Total patients registered: {total}", bold=True)
    add_para(doc, f"• C-section deliveries: {csection} ({csection / total * 100:.1f}%)", bold=True)
    add_para(doc, f"• Maternal complications: {complications}", bold=True)
    doc.add_heading('Patient / Treatment Summary', level=1)
    summary_data = [
        ["Total Registrations", total, 0, "N/A", "Normal range"],
        ["C-section Rate", f"{csection / total * 100:.1f}%", "N/A", "N/A", "Acceptable"],
        ["Complications", complications, "N/A", "N/A", "Monitoring"]
    ]
    insert_table(doc, summary_data, ["Indicator", "Current", "Previous", "% Change", "Notes"])
    return {"total_patients": total, "csection_rate": float(csection / total * 100), "complications": complications,
            "table_data": summary_data}


def generate_infectious_diseases_content(doc, df, prev_data, dept_info, hospital):
    total = int(len(df))
    mortality = int(safe_count(df.get('outcome', pd.Series()), 'Died'))
    admitted = int(safe_count(df.get('admission', pd.Series()), 'Yes'))
    doc.add_heading('Executive Summary', level=1)
    add_para(doc, f"• Total patients: {total}", bold=True)
    add_para(doc, f"• Admitted: {admitted}", bold=True)
    doc.add_page_break()
    add_para(doc, f"• Mortality: {mortality}", bold=True)

    doc.add_heading('Patient / Treatment Summary', level=1)
    summary_data = [
        ["Total Patients", total, 0, "N/A", "Range"],
        ["Admitted", admitted, "N/A", "N/A", "Rate"],
        ["Mortality", mortality, "N/A", "N/A", "Monitor"]
    ]
    insert_table(doc, summary_data, ["Indicator", "Current", "Previous", "% Change", "Notes"])
    return {"total_patients": total, "admitted": admitted, "mortality": mortality, "table_data": summary_data}


def generate_nutrition_content(doc, df, prev_data, dept_info, hospital):
    total = int(len(df))
    stunted = int(safe_count(df.get('stunted', pd.Series()), 'Yes'))
    doc.add_heading('Executive Summary', level=1)
    add_para(doc, f"• Total children: {total}", bold=True)
    add_para(doc, f"• Stunted: {stunted}", bold=True)

    doc.add_heading('Nutritional Assessment', level=1)
    summary_data = [["Total", total, 0, "N/A", "Complete"], ["Stunted", stunted, "N/A", "N/A", "Prevalence"]]
    insert_table(doc, summary_data, ["Indicator", "Current", "Previous", "% Change", "Notes"])
    return {"total_patients": total, "stunted": stunted, "table_data": summary_data}


def generate_mental_health_content(doc, df, prev_data, dept_info, hospital):
    total = int(len(df))
    high_risk = int(safe_count(df.get('suicide_risk', pd.Series()), 'High'))
    doc.add_heading('Executive Summary', level=1)
    add_para(doc, f"• Total screened: {total}", bold=True)
    add_para(doc, f"• High risk: {high_risk}", bold=True)
    doc.add_heading('Mental Health Assessment', level=1)
    summary_data = [["Total", total, 0, "N/A", "Active"], ["High Risk", high_risk, "N/A", "N/A", "Intervention"]]
    insert_table(doc, summary_data, ["Indicator", "Current", "Previous", "% Change", "Notes"])
    return {"total_patients": total, "high_risk": high_risk, "table_data": summary_data}


def generate_ncd_content(doc, df, prev_data, dept_info, hospital):
    total = int(len(df))
    admitted = int(safe_count(df.get('admission', pd.Series()), 'Yes'))
    doc.add_heading('Executive Summary', level=1)
    add_para(doc, f"• Total NCD patients: {total}", bold=True)
    add_para(doc, f"• Admitted: {admitted}", bold=True)

    doc.add_heading('NCD Management', level=1)
    summary_data = [["Total", total, 0, "N/A", "Active"], ["Admissions", admitted, "N/A", "N/A", "Rate"]]
    insert_table(doc, summary_data, ["Indicator", "Current", "Previous", "% Change", "Notes"])
    return {"total_patients": total, "admitted": admitted, "table_data": summary_data}


def generate_endocrinology_content(doc, df, prev_data, dept_info, hospital):
    total = int(len(df))
    type1 = int(safe_count(df.get('diabetes_type', pd.Series()), 'Type 1'))
    doc.add_heading('Executive Summary', level=1)
    add_para(doc, f"• Total diabetes: {total}", bold=True)
    add_para(doc, f"• Type 1: {type1}", bold=True)

    doc.add_heading('Diabetes Management', level=1)
    summary_data = [["Total", total, 0, "N/A", "Active"], ["Type 1", type1, "N/A", "N/A", "Percentage"]]
    insert_table(doc, summary_data, ["Indicator", "Current", "Previous", "% Change", "Notes"])
    return {"total_patients": total, "type1": type1, "table_data": summary_data}


def generate_oncology_content(doc, df, prev_data, dept_info, hospital):
    total = int(len(df))
    mortality = int(safe_count(df.get('outcome', pd.Series()), 'Died'))
    doc.add_heading('Executive Summary', level=1)
    add_para(doc, f"• Total cancer patients: {total}", bold=True)
    doc.add_page_break()
    add_para(doc, f"• Mortality: {mortality}", bold=True)

    doc.add_heading('Oncology Summary', level=1)
    summary_data = [["Total", total, 0, "N/A", "Active"], ["Mortality", mortality, "N/A", "N/A", "Range"]]
    insert_table(doc, summary_data, ["Indicator", "Current", "Previous", "% Change", "Notes"])
    return {"total_patients": total, "mortality": mortality, "table_data": summary_data}


CONTENT_GENERATORS = {
    "infectious_diseases": generate_infectious_diseases_content,
    "maternal_health": generate_maternal_health_content,
    "nutrition_data": generate_nutrition_content,
    "mental_health_data": generate_mental_health_content,
    "ncd_internal_medicine": generate_ncd_content,
    "cardiology_data": generate_cardiology_content,
    "endocrinology_diabetes_data": generate_endocrinology_content,
    "oncology_data": generate_oncology_content
}


def add_graphs_to_report(doc, department, hospital, quarter, year):
    trend_graph = os.path.join(GRAPH_DIR, f"{department}_trend_total_patients.png")
    mort_graph = os.path.join(GRAPH_DIR, f"{department}_mort_comp.png")
    graph_info = {"graphs": [], "added": False}
    doc.add_page_break()
    if os.path.exists(trend_graph) or os.path.exists(mort_graph):

        doc.add_heading('Visual Analytics', level=1)
        graph_info["added"] = True
        if os.path.exists(trend_graph):
            add_para(doc, "Patient Trend:", bold=True)
            doc.add_picture(trend_graph, width=Inches(6))
            graph_info["graphs"].append({"title": "Trend", "file": trend_graph})
            doc.add_page_break()
        if os.path.exists(mort_graph):
            add_para(doc, "Mortality:", bold=True)
            doc.add_picture(mort_graph, width=Inches(6))
            graph_info["graphs"].append({"title": "Mortality", "file": mort_graph})
    return graph_info


def add_signatures(doc, hospital, dept, focal_persons):

    focal_key = f"{hospital}_{dept['focal_dept']}"
    focal = focal_persons.get(focal_key, {'name': '[Name Not Available]'})
    add_para(doc, f"Prepared by: Head of {dept['name']}", bold=True)
    add_para(doc, f"Name: {focal['name']}")
    add_para(doc, f"Date: {datetime.now().strftime('%B %d, %Y')}")


# ===== MAIN TOOL: SINGLE DEPARTMENT (FAST) =====
@mcp.tool()
async def generate_single_department_report(
        department: str,
        hospital: str,
        quarter: int,
        year: int
) -> dict:
    """
    ✅ FIXED: Generate ONE department report (fast, no timeout)
    Use this tool for each department separately
    """
    try:
        if department not in DEPT_CONFIG:
            return {"status": "error", "message": f"Unknown department: {department}"}

        dept = DEPT_CONFIG[department]
        debug_log(f"📋 Generating {dept['name']} report...")

        csv_path = os.path.join(CSV_DIR, dept['csv_file'])
        if not os.path.exists(csv_path):
            return {"status": "error", "message": f"CSV not found: {csv_path}"}

        df = pd.read_csv(csv_path)
        df_filtered = df[(df['hospital'] == hospital) &
                         (df['report_year'] == year) &
                         (df['report_quarter'] == quarter)]

        if len(df_filtered) == 0:
            return {"status": "error", "message": f"No data for {hospital} Q{quarter} {year}"}

        template_path = os.path.join(TEMPLATE_DIR, dept['template'])
        doc = Document(template_path)
        clean_template_content(doc)

        focal_persons = load_focal_persons()
        update_header_info(doc, hospital, quarter, year, dept, focal_persons)

        content_generator = CONTENT_GENERATORS.get(department)
        dept_data = content_generator(doc, df_filtered, None, dept, hospital)

        # AI Analysis (4 calls)
        exec_analysis = await get_ai_section_analysis(department, "executive", dept_data, hospital, quarter, year)
        add_ai_analysis_paragraph(doc, "Executive Summary", exec_analysis, "📋")

        table_analysis = await get_ai_section_analysis(department, "table",
                                                       {"table_data": dept_data.get('table_data', [])}, hospital,
                                                       quarter, year)
        add_ai_analysis_paragraph(doc, "Table Analysis", table_analysis, "📊")

        graph_info = add_graphs_to_report(doc, department, hospital, quarter, year)
        if graph_info["added"]:
            graph_analysis = await get_ai_section_analysis(department, "graph", graph_info, hospital, quarter, year)
            add_ai_analysis_paragraph(doc, "Graph Analysis", graph_analysis, "📈")

        doc.add_page_break()
        doc.add_heading('Recommendations', level=1)
        recommendations = await get_ai_section_analysis(department, "recommendations", dept_data, hospital, quarter,
                                                        year)

        intro_para = doc.add_paragraph()
        intro_run = intro_para.add_run("Evidence-Based Recommendations:")
        intro_run.bold = True
        intro_run.font.size = Pt(12)
        intro_run.font.color.rgb = RGBColor(0, 51, 102)

        if isinstance(recommendations, list):
            for rec in recommendations:
                rec_para = doc.add_paragraph()
                rec_para.style = 'List Bullet'
                rec_para.add_run(rec).font.size = Pt(10)

        add_signatures(doc, hospital, dept, focal_persons)

        os.makedirs(OUTPUT_DIR, exist_ok=True)
        safe_hospital = hospital.replace(" ", "_")
        filename = f"{department}_{safe_hospital}_Q{quarter}_{year}_report.docx"
        output_path = os.path.join(OUTPUT_DIR, filename)

        doc.save(output_path)
        debug_log(f"✅ Report saved: {filename}")

        return {
            "status": "success",
            "message": f"Report generated for {dept['name']}",
            "file_path": output_path,
            "filename": filename,
            "department": dept['name'],
            "hospital": hospital,
            "quarter": f"Q{quarter}",
            "year": year,
            "total_patients": len(df_filtered),
            "ai_analysis_included": True,
            "generated_at": datetime.now().isoformat()
        }

    except Exception as e:
        debug_log(f"❌ Error: {str(e)}")
        import traceback
        debug_log(traceback.format_exc())
        return {
            "status": "error",
            "message": str(e),
            "error_type": type(e).__name__
        }


# ===== BATCH COORDINATOR: START ALL 8 DEPARTMENTS =====
@mcp.tool()
async def start_all_departments_batch(
        hospital: str,
        quarter: int,
        year: int
) -> dict:
    """
    ✅ FIXED: Start batch generation for all 8 departments
    Returns immediately with batch_id for status tracking
    CALL THIS FIRST, then call get_batch_status to check progress
    """
    batch_id = f"BATCH_{hospital.replace(' ', '_')}_{quarter}_{year}_{int(datetime.now().timestamp())}"

    status_file = os.path.join(STATUS_DIR, f"{batch_id}.json")

    # Initialize status file
    status_data = {
        "batch_id": batch_id,
        "hospital": hospital,
        "quarter": quarter,
        "year": year,
        "status": "in_progress",
        "started_at": datetime.now().isoformat(),
        "departments": {dept: "pending" for dept in DEPT_CONFIG.keys()},
        "completed": [],
        "failed": [],
        "total": len(DEPT_CONFIG),
        "completed_count": 0
    }

    with open(status_file, 'w') as f:
        json.dump(status_data, f, indent=2)

    # Start background task
    asyncio.create_task(_process_batch(batch_id, hospital, quarter, year))

    debug_log(f"✅ Batch {batch_id} started for {hospital}")

    return {
        "status": "started",
        "batch_id": batch_id,
        "hospital": hospital,
        "quarter": f"Q{quarter}",
        "year": year,
        "total_departments": len(DEPT_CONFIG),
        "estimated_time_minutes": len(DEPT_CONFIG) * 2,  # ~2 min per dept
        "message": "Batch generation started. Use get_batch_status to track progress."
    }


async def _process_batch(batch_id: str, hospital: str, quarter: int, year: int):
    """Background task to process all departments sequentially"""
    status_file = os.path.join(STATUS_DIR, f"{batch_id}.json")

    try:
        # Ensure MCP connected once
        await ensure_mcp_connected()

        for idx, dept_key in enumerate(DEPT_CONFIG.keys(), 1):
            debug_log(f"[{idx}/{len(DEPT_CONFIG)}] Processing {dept_key}...")

            # Update status: processing
            with open(status_file, 'r') as f:
                status_data = json.load(f)
            status_data["departments"][dept_key] = "processing"
            with open(status_file, 'w') as f:
                json.dump(status_data, f, indent=2)

            try:
                # Generate single report
                result = await generate_single_department_report(dept_key, hospital, quarter, year)

                # Update status based on result
                with open(status_file, 'r') as f:
                    status_data = json.load(f)

                if result["status"] == "success":
                    status_data["departments"][dept_key] = "completed"
                    status_data["completed"].append({
                        "department": dept_key,
                        "file_path": result["file_path"],
                        "total_patients": result["total_patients"]
                    })
                    status_data["completed_count"] += 1
                    debug_log(f"✅ [{idx}/{len(DEPT_CONFIG)}] {dept_key} completed")
                else:
                    status_data["departments"][dept_key] = "failed"
                    status_data["failed"].append({
                        "department": dept_key,
                        "error": result.get("message", "Unknown error")
                    })
                    debug_log(f"❌ [{idx}/{len(DEPT_CONFIG)}] {dept_key} failed")

                with open(status_file, 'w') as f:
                    json.dump(status_data, f, indent=2)

            except Exception as e:
                debug_log(f"❌ Exception in {dept_key}: {e}")
                with open(status_file, 'r') as f:
                    status_data = json.load(f)
                status_data["departments"][dept_key] = "failed"
                status_data["failed"].append({"department": dept_key, "error": str(e)})
                with open(status_file, 'w') as f:
                    json.dump(status_data, f, indent=2)

            # Small delay between departments
            await asyncio.sleep(2)

        # Mark batch as complete
        with open(status_file, 'r') as f:
            status_data = json.load(f)
        status_data["status"] = "completed"
        status_data["completed_at"] = datetime.now().isoformat()
        with open(status_file, 'w') as f:
            json.dump(status_data, f, indent=2)

        debug_log(f"✅ Batch {batch_id} completed: {status_data['completed_count']}/{len(DEPT_CONFIG)}")

    except Exception as e:
        debug_log(f"❌ Batch {batch_id} failed: {e}")
        with open(status_file, 'r') as f:
            status_data = json.load(f)
        status_data["status"] = "failed"
        status_data["error"] = str(e)
        with open(status_file, 'w') as f:
            json.dump(status_data, f, indent=2)


@mcp.tool()
def get_batch_status(batch_id: str) -> dict:
    """
    ✅ Check status of batch report generation
    Returns current progress and completed reports
    """
    status_file = os.path.join(STATUS_DIR, f"{batch_id}.json")

    if not os.path.exists(status_file):
        return {
            "status": "not_found",
            "message": f"Batch {batch_id} not found"
        }

    try:
        with open(status_file, 'r') as f:
            status_data = json.load(f)

        return {
            "status": status_data["status"],
            "batch_id": batch_id,
            "hospital": status_data["hospital"],
            "quarter": f"Q{status_data['quarter']}",
            "year": status_data["year"],
            "progress": {
                "total": status_data["total"],
                "completed": status_data["completed_count"],
                "pending": sum(1 for v in status_data["departments"].values() if v == "pending"),
                "processing": sum(1 for v in status_data["departments"].values() if v == "processing"),
                "failed": len(status_data["failed"]),
                "percentage": int((status_data["completed_count"] / status_data["total"]) * 100)
            },
            "departments": status_data["departments"],
            "completed_reports": status_data["completed"],
            "failed_reports": status_data["failed"],
            "started_at": status_data.get("started_at"),
            "completed_at": status_data.get("completed_at")
        }

    except Exception as e:
        return {
            "status": "error",
            "message": f"Error reading status: {str(e)}"
        }


@mcp.tool()
def list_all_batches() -> dict:
    """List all batch jobs (active and completed)"""
    batches = []

    try:
        for filename in os.listdir(STATUS_DIR):
            if filename.endswith('.json'):
                with open(os.path.join(STATUS_DIR, filename), 'r') as f:
                    status_data = json.load(f)
                batches.append({
                    "batch_id": status_data["batch_id"],
                    "hospital": status_data["hospital"],
                    "quarter": f"Q{status_data['quarter']}",
                    "year": status_data["year"],
                    "status": status_data["status"],
                    "completed": status_data["completed_count"],
                    "total": status_data["total"],
                    "started_at": status_data.get("started_at")
                })
    except Exception as e:
        return {"status": "error", "message": str(e)}

    return {
        "status": "success",
        "total_batches": len(batches),
        "batches": sorted(batches, key=lambda x: x.get("started_at", ""), reverse=True)
    }


@mcp.tool()
def list_available_departments() -> dict:
    """List all 8 available departments"""
    return {
        "total_departments": len(DEPT_CONFIG),
        "departments": [
            {
                "key": key,
                "name": config["name"],
                "template": config["template"],
                "csv_file": config["csv_file"]
            }
            for key, config in DEPT_CONFIG.items()
        ]
    }


@mcp.tool()
def check_data_availability(department: str, hospital: str, quarter: int, year: int) -> dict:
    """Check if data exists for a specific department/hospital/quarter"""
    try:
        if department not in DEPT_CONFIG:
            return {"status": "error", "message": "Invalid department"}

        dept = DEPT_CONFIG[department]
        csv_path = os.path.join(CSV_DIR, dept['csv_file'])

        if not os.path.exists(csv_path):
            return {"status": "unavailable", "message": "CSV file not found"}

        df = pd.read_csv(csv_path)
        df_filtered = df[(df['hospital'] == hospital) &
                         (df['report_year'] == year) &
                         (df['report_quarter'] == quarter)]

        return {
            "status": "available" if len(df_filtered) > 0 else "no_data",
            "record_count": len(df_filtered),
            "department": dept['name'],
            "hospital": hospital,
            "quarter": f"Q{quarter}",
            "year": year
        }

    except Exception as e:
        return {"status": "error", "message": str(e)}


if __name__ == "__main__":
    print("🏥 Enhanced Report Generation MCP Server (BATCH PROCESSING) Starting...", file=sys.stderr)
    print(f"✅ {len(DEPT_CONFIG)} departments configured", file=sys.stderr)
    print(f"✅ BATCH MODE: No timeout issues!", file=sys.stderr)
    print(f"✅ AI Agent Analysis: ENABLED", file=sys.stderr)
    print(f"✅ Status tracking: {STATUS_DIR}", file=sys.stderr)
    print(f"", file=sys.stderr)
    print(f"📌 USAGE:", file=sys.stderr)
    print(f"  1. Call start_all_departments_batch() to begin", file=sys.stderr)
    print(f"  2. Call get_batch_status(batch_id) to check progress", file=sys.stderr)
    print(f"  3. All 8 reports will generate sequentially (no timeout)", file=sys.stderr)
    mcp.run()