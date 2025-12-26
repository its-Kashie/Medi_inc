# rnd_mcp_tools.py - R&D Agent MCP Tools
"""
R&D-specific tools for:
1. University email campaigns with focal persons
2. WHO proposal generation (Word documents)
3. Research priority analysis
4. 3-year trend visualization
"""

import os
import sys
import json
import time
from datetime import datetime
from functools import wraps
from email.message import EmailMessage
import smtplib
import threading
import requests

from mcp.server.fastmcp import FastMCP

os.environ['MCP_CLIENT_TIMEOUT'] = '10'


# ===== DEBUG LOGGING =====
def debug_log(msg):
    print(msg, file=sys.stderr, flush=True)


debug_log("✅ R&D MCP Tools initialized")

# ===== MCP SERVER =====
mcp = FastMCP("R&D Research Collaboration Tools")


# ===== LAZY IMPORTS =====
def get_mongo():
    global MongoClient, mongo_client, db
    if 'mongo_client' not in globals():
        try:
            from pymongo import MongoClient
            mongo_client = MongoClient("mongodb://localhost:27017/", serverSelectionTimeoutMS=3000)
            db = mongo_client["healthcare360_full"]
            mongo_client.admin.command('ping')
            debug_log("✅ MongoDB connected (R&D MCP)")
            return mongo_client, db
        except Exception as e:
            debug_log(f"⚠️ MongoDB unavailable: {e}")
            mongo_client = None
            db = None
            return None, None
    return mongo_client, db


def get_docx():
    global Document, Inches, Pt, RGBColor, WD_ALIGN_PARAGRAPH
    if 'Document' not in globals():
        from docx import Document
        from docx.shared import Inches, Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    return Document, Inches, Pt, RGBColor, WD_ALIGN_PARAGRAPH


def get_plotting():
    global pd, plt, matplotlib, io
    if 'pd' not in globals():
        import pandas as pd
        import matplotlib
        matplotlib.use('Agg')
        import matplotlib.pyplot as plt
        import io
    return pd, plt, matplotlib, io


 #===== ADD BROADCAST FUNCTION =====
def broadcast_trace_to_dashboard(action: str, details: dict):
    """Broadcast trace to dashboard via API server"""
    try:
        requests.post(
            'http://localhost:8000/api/internal/broadcast-trace',
            json={
                "agent": "rnd",
                "action": action,
                "details": details
            },
            timeout=2
        )
    except Exception as e:
        debug_log(f"⚠️ Failed to broadcast trace: {e}")


# ===== DIRECTORIES =====
WHO_PROPOSALS_DIR = "generated_reports/who_proposals"
GRAPHS_DIR = "generated_reports/graphs"
os.makedirs(WHO_PROPOSALS_DIR, exist_ok=True)
os.makedirs(GRAPHS_DIR, exist_ok=True)

# ===== EMAIL CONFIG =====
# 'nooreasal786@gmail.com'
# "irph tole tuqr vfmi "
RESEARCH_SENDER_EMAIL = "c.healthcare360@gmail.com"
APP_PASSWORD =   "djhi xaty geuj pvyy"

# ===== TRACE LOG =====
TRACE_LOG = []


def log_trace(action: str, details: dict):
    entry = {
        "timestamp": datetime.now().isoformat(),
        "action": action,
        "details": details
    }
    TRACE_LOG.append(entry)
    debug_log(f"[R&D-MCP] {action}: {details}")


def traced_tool(func):
    @wraps(func)
    def wrapper(*args, **kwargs):
        call_id = f"{func.__name__}-{int(time.time() * 1000)}"
        log_trace("tool_start", {"tool": func.__name__, "call_id": call_id})
        try:
            result = func(*args, **kwargs)
            log_trace("tool_success", {"tool": func.__name__, "call_id": call_id})
            return result
        except Exception as e:
            log_trace("tool_error", {"tool": func.__name__, "error": str(e)})
            raise

    return wrapper


# ===== DEPARTMENT MAPPING =====
DEPT_KEY_MAPPING = {
    "Cardiology": "cardiology_data",
    "Maternal Health": "maternal_health",
    "Infectious Diseases": "infectious_diseases",
    "Nutrition": "nutrition_data",
    "Mental Health": "mental_health_data",
    "NCD": "ncd_internal_medicine",
    "Endocrinology": "endocrinology_diabetes_data",
    "Oncology": "oncology_data"
}


# =============================================
# FOCAL PERSONS MANAGEMENT
# =============================================

@mcp.tool()
@traced_tool
def get_university_focal_persons() -> dict:
    """
    Load university focal persons from Excel file.
    Returns list of all universities with contact information.

    Returns:
        dict with focal persons list and metadata
    """
    try:
        pd, _, _, _ = get_plotting()

        excel_path = "focal_persons_excels/university_focal_persons.xlsx"

        if not os.path.exists(excel_path):
            return {
                "status": "error",
                "message": f"Focal persons Excel file not found at {excel_path}",
                "focal_persons": []
            }

        df = pd.read_excel(excel_path)

        focal_persons = []
        for _, row in df.iterrows():
            focal_persons.append({
                "university": str(row['University']),
                "department": str(row['Department']),
                "focal_person_name": str(row['Focal Person Name']),
                "email": str(row['Email']),
                "contact": str(row['Contact']),
                "notes": str(row.get('Notes / Internship Opportunities', ''))
            })

        log_trace("focal_persons_loaded", {
            "total_universities": len(focal_persons),
            "source": excel_path
        })

        return {
            "status": "success",
            "total_universities": len(focal_persons),
            "focal_persons": focal_persons,
            "loaded_from": excel_path
        }

    except Exception as e:
        log_trace("focal_persons_load_error", {"error": str(e)})
        return {
            "status": "error",
            "message": str(e),
            "focal_persons": []
        }


# =============================================
# UNIVERSITY EMAIL CAMPAIGNS
# =============================================

# @mcp.tool()
# @traced_tool
# def send_university_collaboration_emails(
#         research_area: str,
#         evidence_summary: str,
#         internship_count: int = 10,
#         attachment_paths: list = None
# ) -> dict:
#     """
#     Send research collaboration emails to ALL universities from Excel.
#     Uses threading for parallel sending.
#
#     Args:
#         research_area: Research topic (e.g., "Maternal Health Crisis")
#         evidence_summary: Summary of findings from national data
#         internship_count: Number of internship positions available
#         attachment_paths: List of file paths to attach (graphs, reports)
#
#     Returns:
#         dict with email sending results
#     """
#
#     # Load focal persons
#     focal_data = get_university_focal_persons()
#
#     if focal_data["status"] != "success":
#         return focal_data
#
#     focal_persons = focal_data["focal_persons"]
#
#     if not focal_persons:
#         return {
#             "status": "error",
#             "message": "No focal persons found in Excel file"
#         }
#
#     results = []
#
#     def send_email_thread(focal_person):
#         try:
#             msg = EmailMessage()
#             msg['Subject'] = f"Healthcare Research Collaboration: {research_area}"
#             msg['From'] = RESEARCH_SENDER_EMAIL
#             msg['To'] = focal_person["email"]
#
#             # Email body
#             email_body = f"""
# Dear {focal_person['focal_person_name']},
#
# The National Institute of Health (NIH) Pakistan has identified {research_area} as a HIGH PRIORITY research area based on analysis of 10 major hospitals over 3 years (2023-2025).
#
# **Evidence Summary:**
# {evidence_summary}
#
# **Collaboration Opportunities for {focal_person['university']} - {focal_person['department']}:**
#
# 1. **Student Internships:** {internship_count} positions available at partner hospitals
# 2. **Final Year Projects (FYP):** Access to real healthcare data and mentorship
# 3. **Research Partnerships:** WHO-funded initiatives with international visibility
# 4. **Training Programs:** Skill development for students in data analysis and healthcare research
#
# **What We Offer:**
# - Access to 3 years of anonymized healthcare data from 10 hospitals
# - Mentorship from NIH researchers and hospital clinicians
# - Publication opportunities in international journals
# - Research funding support through WHO partnerships
# - Certificate of completion for internships
#
# **Next Steps:**
# 1. Review attached data analysis and research priorities
# 2. Schedule a meeting with NIH R&D team to discuss collaboration
# 3. Identify interested students/faculty for placement
# 4. Sign MOU for data sharing and research collaboration
#
# **Contact Information:**
# Research & Development Division
# National Institute of Health Pakistan
# Email: {RESEARCH_SENDER_EMAIL}
#
# We look forward to partnering with {focal_person['university']} to advance healthcare research in Pakistan.
#
# Best regards,
# R&D Division
# National Institute of Health Pakistan
# """
#
#             msg.set_content(email_body)
#
#             # Add attachments
#             if attachment_paths:
#                 for attachment_path in attachment_paths:
#                     if os.path.exists(attachment_path):
#                         with open(attachment_path, 'rb') as f:
#                             file_data = f.read()
#                             file_name = os.path.basename(attachment_path)
#
#                             # Determine MIME type
#                             if file_name.endswith('.pdf'):
#                                 maintype, subtype = 'application', 'pdf'
#                             elif file_name.endswith('.docx'):
#                                 maintype, subtype = 'application', 'vnd.openxmlformats-officedocument.wordprocessingml.document'
#                             elif file_name.endswith('.png'):
#                                 maintype, subtype = 'image', 'png'
#                             else:
#                                 maintype, subtype = 'application', 'octet-stream'
#
#                             msg.add_attachment(file_data, maintype=maintype, subtype=subtype, filename=file_name)
#
#             # Send email
#             with smtplib.SMTP("smtp.gmail.com", 587, timeout=10) as smtp:
#                 smtp.starttls()
#                 smtp.login(RESEARCH_SENDER_EMAIL, APP_PASSWORD)
#                 smtp.send_message(msg)
#
#             results.append({
#                 "university": focal_person["university"],
#                 "focal_person": focal_person["focal_person_name"],
#                 "email": focal_person["email"],
#                 "status": "sent",
#                 "sent_at": datetime.now().isoformat()
#             })
#
#         except Exception as e:
#             results.append({
#                 "university": focal_person["university"],
#                 "focal_person": focal_person["focal_person_name"],
#                 "email": focal_person["email"],
#                 "status": "failed",
#                 "error": str(e)
#             })
#
#     # Send emails in parallel using threads
#     threads = [threading.Thread(target=send_email_thread, args=(fp,)) for fp in focal_persons]
#
#     for t in threads:
#         t.start()
#
#     for t in threads:
#         t.join()
#
#     # Calculate statistics
#     sent_count = len([r for r in results if r['status'] == 'sent'])
#     failed_count = len([r for r in results if r['status'] == 'failed'])
#
#     log_trace("university_emails_sent", {
#         "research_area": research_area,
#         "total": len(focal_persons),
#         "sent": sent_count,
#         "failed": failed_count
#     })
#
#     return {
#         "status": "completed",
#         "research_area": research_area,
#         "total_universities": len(focal_persons),
#         "emails_sent": sent_count,
#         "emails_failed": failed_count,
#         "success_rate": f"{(sent_count / len(focal_persons) * 100):.1f}%",
#         "details": results,
#         "campaign_completed_at": datetime.now().isoformat()
#     }
#


# ===== UPDATE send_university_collaboration_emails =====
@mcp.tool()
@traced_tool
def send_university_collaboration_emails(
        research_area: str,
        evidence_summary: str,
        internship_count: int = 10,
        attachment_paths: list = None
) -> dict:
    """
    Send research collaboration emails to ALL universities from Excel.
    Uses threading for parallel sending.
    """

    # Broadcast START
    broadcast_trace_to_dashboard("university_email_campaign_started", {
        "research_area": research_area,
        "status": "starting"
    })

    # Load focal persons
    focal_data = get_university_focal_persons()

    if focal_data["status"] != "success":
        return focal_data

    focal_persons = focal_data["focal_persons"]

    if not focal_persons:
        return {
            "status": "error",
            "message": "No focal persons found in Excel file"
        }

    # Broadcast LOADED
    broadcast_trace_to_dashboard("universities_loaded", {
        "total_universities": len(focal_persons),
        "research_area": research_area
    })

    results = []

    def send_email_thread(focal_person):
        try:
            msg = EmailMessage()
            msg['Subject'] = f"Healthcare Research Collaboration: {research_area}"
            msg['From'] = RESEARCH_SENDER_EMAIL
            msg['To'] = focal_person["email"]

            # Email body (keep same as before)
            email_body = f"""
Dear {focal_person['focal_person_name']},

The National Institute of Health (NIH) Pakistan has identified {research_area} as a HIGH PRIORITY research area based on analysis of 10 major hospitals over 3 years (2023-2025).

**Evidence Summary:**
{evidence_summary}

**Collaboration Opportunities for {focal_person['university']} - {focal_person['department']}:**

1. **Student Internships:** {internship_count} positions available at partner hospitals
2. **Final Year Projects (FYP):** Access to real healthcare data and mentorship
3. **Research Partnerships:** WHO-funded initiatives with international visibility
4. **Training Programs:** Skill development for students in data analysis and healthcare research

**What We Offer:**
- Access to 3 years of anonymized healthcare data from 10 hospitals
- Mentorship from NIH researchers and hospital clinicians
- Publication opportunities in international journals
- Research funding support through WHO partnerships
- Certificate of completion for internships

**Next Steps:**
1. Review attached data analysis and research priorities
2. Schedule a meeting with NIH R&D team to discuss collaboration
3. Identify interested students/faculty for placement
4. Sign MOU for data sharing and research collaboration

**Contact Information:**
Research & Development Division
National Institute of Health Pakistan
Email: {RESEARCH_SENDER_EMAIL}

We look forward to partnering with {focal_person['university']} to advance healthcare research in Pakistan.

Best regards,
R&D Division
National Institute of Health Pakistan
"""

            msg.set_content(email_body)

            # Add attachments (if any)
            if attachment_paths:
                for attachment_path in attachment_paths:
                    if os.path.exists(attachment_path):
                        with open(attachment_path, 'rb') as f:
                            file_data = f.read()
                            file_name = os.path.basename(attachment_path)

                            if file_name.endswith('.pdf'):
                                maintype, subtype = 'application', 'pdf'
                            elif file_name.endswith('.docx'):
                                maintype, subtype = 'application', 'vnd.openxmlformats-officedocument.wordprocessingml.document'
                            elif file_name.endswith('.png'):
                                maintype, subtype = 'image', 'png'
                            else:
                                maintype, subtype = 'application', 'octet-stream'

                            msg.add_attachment(file_data, maintype=maintype, subtype=subtype, filename=file_name)

            # Send email
            with smtplib.SMTP("smtp.gmail.com", 587, timeout=10) as smtp:
                smtp.starttls()
                smtp.login(RESEARCH_SENDER_EMAIL, APP_PASSWORD)
                smtp.send_message(msg)

            results.append({
                "university": focal_person["university"],
                "focal_person": focal_person["focal_person_name"],
                "email": focal_person["email"],
                "status": "sent",
                "sent_at": datetime.now().isoformat()
            })

            # Broadcast EACH successful send
            broadcast_trace_to_dashboard("email_sent", {
                "university": focal_person["university"],
                "email": focal_person["email"],
                "research_area": research_area
            })

        except Exception as e:
            results.append({
                "university": focal_person["university"],
                "focal_person": focal_person["focal_person_name"],
                "email": focal_person["email"],
                "status": "failed",
                "error": str(e)
            })

            # Broadcast FAILURE
            broadcast_trace_to_dashboard("email_failed", {
                "university": focal_person["university"],
                "error": str(e)[:100]
            })

    # Send emails in parallel using threads
    threads = [threading.Thread(target=send_email_thread, args=(fp,)) for fp in focal_persons]

    for t in threads:
        t.start()

    for t in threads:
        t.join()

    # Calculate statistics
    sent_count = len([r for r in results if r['status'] == 'sent'])
    failed_count = len([r for r in results if r['status'] == 'failed'])

    # Broadcast FINAL RESULTS
    broadcast_trace_to_dashboard("university_email_campaign_completed", {
        "research_area": research_area,
        "total_universities": len(focal_persons),
        "emails_sent": sent_count,
        "emails_failed": failed_count,
        "success_rate": f"{(sent_count / len(focal_persons) * 100):.1f}%"
    })

    log_trace("university_emails_sent", {
        "research_area": research_area,
        "total": len(focal_persons),
        "sent": sent_count,
        "failed": failed_count
    })

    return {
        "status": "completed",
        "research_area": research_area,
        "total_universities": len(focal_persons),
        "emails_sent": sent_count,
        "emails_failed": failed_count,
        "success_rate": f"{(sent_count / len(focal_persons) * 100):.1f}%",
        "details": results,
        "campaign_completed_at": datetime.now().isoformat()
    }
# =============================================
# WHO PROPOSAL GENERATION (WORD DOCUMENTS)
# =============================================

@mcp.tool()
@traced_tool
def generate_who_funding_proposal_docx(
        research_area: str,
        national_data_summary: dict,
        three_year_trends: dict = None,
        priority_justification: str = None
) -> dict:
    """
    Generate comprehensive WHO funding proposal as Word document.
    Includes graphs, evidence, and complete proposal structure.

    Args:
        research_area: Research focus (e.g., "Maternal Health Crisis in Pakistan")
        national_data_summary: Dict with aggregated national statistics
        three_year_trends: Optional 3-year trend data for graphs
        priority_justification: Why this is high priority

    Returns:
        dict with file path and proposal metadata
    """

    try:
        Document, Inches, Pt, RGBColor, WD_ALIGN_PARAGRAPH = get_docx()
        pd, plt, matplotlib, io = get_plotting()

        # Create document
        doc = Document()

        # ===== TITLE PAGE =====
        title = doc.add_heading('WHO FUNDING PROPOSAL', level=0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        subtitle = doc.add_heading(f'Indigenous Health Research: {research_area}', level=1)
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

        doc.add_paragraph(f"\nSubmitted by: National Institute of Health (NIH) Pakistan")
        doc.add_paragraph(f"Date: {datetime.now().strftime('%B %d, %Y')}")
        doc.add_paragraph(f"Proposal ID: WHO-PROP-{int(time.time())}")

        doc.add_page_break()

        # ===== EXECUTIVE SUMMARY =====
        doc.add_heading('Executive Summary', level=1)

        exec_para = doc.add_paragraph()
        exec_para.add_run("Problem Statement: ").bold = True
        exec_para.add_run(
            f"National healthcare data from 10 major hospitals reveals critical gaps in {research_area}. "
            f"Analysis of {national_data_summary.get('total_patients_all_departments', 0)} patients over 3 years "
            f"(2023-2025) indicates urgent need for intervention."
        )

        if priority_justification:
            doc.add_paragraph()
            priority_para = doc.add_paragraph()
            priority_para.add_run("Priority Justification: ").bold = True
            priority_para.add_run(priority_justification)

        doc.add_page_break()

        # ===== DISEASE BURDEN ANALYSIS =====
        doc.add_heading('National Disease Burden Analysis', level=1)

        # Add summary table
        if 'department_breakdown' in national_data_summary:
            table = doc.add_table(rows=1, cols=4)
            table.style = 'Light Grid Accent 1'

            headers = ['Department', 'Total Cases (National)', 'Hospitals Reporting', 'Avg per Hospital']
            for i, header in enumerate(headers):
                cell = table.rows[0].cells[i]
                cell.text = header
                cell.paragraphs[0].runs[0].font.bold = True

            # Add department rows
            for dept_name, dept_data in national_data_summary.get('department_breakdown', {}).items():
                row = table.add_row()
                row.cells[0].text = dept_name
                row.cells[1].text = str(dept_data.get('total_cases_national', 0))
                row.cells[2].text = str(dept_data.get('hospitals_reporting', 0))
                row.cells[3].text = str(dept_data.get('average_per_hospital', 0))

        doc.add_page_break()

        # ===== 3-YEAR TREND GRAPHS =====
        if three_year_trends and 'quarterly_breakdown' in three_year_trends:
            doc.add_heading('Three-Year Trend Analysis (2023-2025)', level=1)

            quarters = [d['period'] for d in three_year_trends['quarterly_breakdown']]
            patients = [d['total_patients'] for d in three_year_trends['quarterly_breakdown']]

            plt.figure(figsize=(10, 6))
            plt.plot(quarters, patients, marker='o', linewidth=2, markersize=8, color='#2563eb')
            plt.title(f"Patient Trends: {research_area} (2023-2025)", fontsize=14, fontweight='bold')
            plt.xlabel('Quarter', fontsize=12)
            plt.ylabel('Total Patients', fontsize=12)
            plt.xticks(rotation=45, ha='right')
            plt.grid(True, alpha=0.3)
            plt.tight_layout()

            # Save graph
            graph_path = os.path.join(GRAPHS_DIR, f"trend_{int(time.time())}.png")
            plt.savefig(graph_path, dpi=300, bbox_inches='tight')
            plt.close()

            # Add to document
            doc.add_picture(graph_path, width=Inches(6))

            doc.add_paragraph()
            analysis = doc.add_paragraph()
            analysis.add_run("Key Finding: ").bold = True
            analysis.add_run(
                f"Total patients analyzed over 3 years: {three_year_trends.get('total_patients_3yr', 0)}. "
                f"Average per quarter: {three_year_trends.get('avg_per_quarter', 0):.0f} patients."
            )

        doc.add_page_break()

        # ===== FUNDING REQUEST =====
        doc.add_heading('Funding Request', level=1)

        funding_table = doc.add_table(rows=6, cols=2)
        funding_table.style = 'Light Grid Accent 1'

        funding_data = [
            ['Total Amount Requested', '$500,000 USD'],
            ['Project Duration', '24 months'],
            ['Research Staff', '$200,000'],
            ['Equipment & Infrastructure', '$150,000'],
            ['Training & Capacity Building', '$100,000'],
            ['Administrative Overhead', '$50,000']
        ]

        for i, (item, amount) in enumerate(funding_data):
            funding_table.rows[i].cells[0].text = item
            funding_table.rows[i].cells[1].text = amount
            if i == 0:
                funding_table.rows[i].cells[0].paragraphs[0].runs[0].font.bold = True
                funding_table.rows[i].cells[1].paragraphs[0].runs[0].font.bold = True

        doc.add_page_break()

        # ===== INTERVENTION PLAN =====
        doc.add_heading('Proposed Interventions', level=1)

        interventions = [
            f"Expand {research_area.split()[0]} screening programs across all 10 hospitals",
            "Train 100+ healthcare workers in evidence-based protocols",
            "Develop culturally-appropriate intervention models for Pakistan",
            "Establish real-time monitoring and reporting systems",
            "Create community awareness campaigns",
            "Implement quality improvement initiatives"
        ]

        for intervention in interventions:
            p = doc.add_paragraph(intervention, style='List Bullet')
            p.paragraph_format.space_after = Pt(6)

        doc.add_page_break()

        # ===== EXPECTED OUTCOMES =====
        doc.add_heading('Expected Outcomes', level=1)

        outcomes = [
            "Reduce mortality rates by 20% in target population",
            "Improve early detection rates by 30%",
            "Train 100+ healthcare professionals",
            "Publish findings in 3+ international peer-reviewed journals",
            "Create sustainable, locally-owned solutions",
            "Establish national disease registry",
            "Develop policy recommendations for Ministry of Health"
        ]

        for outcome in outcomes:
            p = doc.add_paragraph(outcome, style='List Bullet')
            p.paragraph_format.space_after = Pt(6)

        doc.add_page_break()

        # ===== IMPLEMENTATION PARTNERS =====
        doc.add_heading('Implementation Partners', level=1)

        partners = [
            "National Institute of Health (NIH) Pakistan - Lead Agency",
            "10 Major Teaching Hospitals in Lahore",
            "King Edward Medical University",
            "University of Health Sciences Lahore",
            "Fatima Jinnah Medical University",
            "Pakistan Medical Research Council",
            "World Health Organization (WHO) Pakistan Office"
        ]

        for partner in partners:
            doc.add_paragraph(partner, style='List Bullet')

        doc.add_page_break()

        # ===== SUSTAINABILITY =====
        doc.add_heading('Sustainability Plan', level=1)

        doc.add_paragraph(
            "This project will establish permanent infrastructure and trained personnel "
            "within Pakistan's healthcare system. Government buy-in and local ownership "
            "ensure long-term sustainability beyond WHO funding period. All interventions "
            "are designed to be integrated into existing healthcare delivery systems."
        )

        # ===== SIGNATURES =====
        doc.add_page_break()
        doc.add_heading('Approval & Signatures', level=1)

        doc.add_paragraph(f"\n\n_______________________________")
        doc.add_paragraph("Director General, NIH Pakistan")
        doc.add_paragraph(f"Date: {datetime.now().strftime('%B %d, %Y')}")

        doc.add_paragraph(f"\n\n_______________________________")
        doc.add_paragraph("Secretary, Ministry of Health")
        doc.add_paragraph(f"Date: {datetime.now().strftime('%B %d, %Y')}")

        # ===== SAVE DOCUMENT =====
        safe_research_area = research_area.replace(" ", "_").replace("/", "-")
        filename = f"WHO_Proposal_{safe_research_area}_{datetime.now().strftime('%Y%m%d')}.docx"
        output_path = os.path.join(WHO_PROPOSALS_DIR, filename)

        doc.save(output_path)

        file_size_mb = round(os.path.getsize(output_path) / (1024 * 1024), 2)

        log_trace("who_proposal_generated", {
            "research_area": research_area,
            "file_path": output_path,
            "file_size_mb": file_size_mb
        })

        return {
            "status": "success",
            "message": "WHO funding proposal generated successfully",
            "file_path": output_path,
            "filename": filename,
            "download_url": f"/api/proposals/download?file={filename}",
            "research_area": research_area,
            "total_hospitals": national_data_summary.get('total_hospitals', 10),
            "total_patients_analyzed": national_data_summary.get('total_patients_all_departments', 0),
            "funding_requested_usd": 500000,
            "generated_at": datetime.now().isoformat(),
            "file_size_mb": file_size_mb
        }

    except Exception as e:
        log_trace("who_proposal_error", {"error": str(e)})
        return {
            "status": "error",
            "message": f"Failed to generate WHO proposal: {str(e)}"
        }


# ===== INTERNSHIP DOCUMENT GENERATOR =====
@mcp.tool()
@traced_tool
def generate_internship_call_document(
        research_area: str,
        evidence_summary: str,
        internship_count: int = 10,
        duration_months: int = 6
) -> dict:
    """
    Generate filled NIH Internship Call document for a specific research area.
    Fills all blanks and adds research-specific details.

    Args:
        research_area: Research topic (e.g., "Maternal Health Crisis")
        evidence_summary: Evidence from national data
        internship_count: Number of positions
        duration_months: Internship duration in months

    Returns:
        dict with file path of generated document
    """
    try:
        Document, Inches, Pt, RGBColor, WD_ALIGN_PARAGRAPH = get_docx()
        from datetime import datetime, timedelta

        # Load template
        template_path = "focal_persons_excels/National Health Internship Call(Research_agent_to_university).docx"

        if not os.path.exists(template_path):
            return {
                "status": "error",
                "message": f"Template not found at {template_path}"
            }

        doc = Document(template_path)

        # ===== CALCULATE DATES =====
        today = datetime.now()
        deadline = today + timedelta(days=30)  # 30 days from today
        shortlist_date = deadline + timedelta(days=7)  # 7 days after deadline
        start_date = shortlist_date + timedelta(days=14)  # 2 weeks after shortlisting
        end_date = start_date + timedelta(days=duration_months * 30)  # Based on duration

        # ===== CONTACT INFO (Assumptions) =====
        contact_info = {
            "address": "National Institute of Health, Park Road, Chak Shahzad, Islamabad, Pakistan",
            "email": "researchdivision@nih.gov.pk",
            "phone": "+92-51-9255590",
            "website": "www.nih.gov.pk"
        }

        # ===== FILL IN THE BLANKS =====
        # We'll replace all _______ with actual values

        replacements = [
            duration_months,  # 1st occurrence: Duration
            deadline.strftime("%B %d, %Y"),  # 2nd: Application Deadline
            contact_info["email"],  # 3rd: Submission Email
            contact_info["email"],  # 4th: Submission Email (in Section 6)
            deadline.strftime("%B %d, %Y"),  # 5th: Deadline (in Section 6)
            today.strftime("%B %d, %Y"),  # 6th: Release Date (table)
            deadline.strftime("%B %d, %Y"),  # 7th: Deadline (table)
            shortlist_date.strftime("%B %d, %Y"),  # 8th: Shortlisting Date
            start_date.strftime("%B %d, %Y"),  # 9th: Start Date
            end_date.strftime("%B %d, %Y")  # 10th: End Date
        ]

        # Replace in paragraphs
        occurrence_count = 0
        for paragraph in doc.paragraphs:
            while "_______" in paragraph.text and occurrence_count < len(replacements):
                paragraph.text = paragraph.text.replace("_______", str(replacements[occurrence_count]), 1)
                occurrence_count += 1

        # Replace in tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    while "_______" in cell.text and occurrence_count < len(replacements):
                        cell.text = cell.text.replace("_______", str(replacements[occurrence_count]), 1)
                        occurrence_count += 1

        # ===== REPLACE CONTACT INFO PLACEHOLDERS =====
        for paragraph in doc.paragraphs:
            # Replace [Address]
            if "[Address]" in paragraph.text:
                paragraph.text = paragraph.text.replace("[Address]", contact_info["address"])

            # Replace [Email]
            if "[Email]" in paragraph.text:
                paragraph.text = paragraph.text.replace("[Email]", contact_info["email"])

            # Replace [Phone]
            if "[Phone]" in paragraph.text:
                paragraph.text = paragraph.text.replace("[Phone]", contact_info["phone"])

            # Replace [www.nih.gov.pk]
            if "[www.nih.gov.pk]" in paragraph.text:
                paragraph.text = paragraph.text.replace("[www.nih.gov.pk]", contact_info["website"])

        # ===== ADD RESEARCH-SPECIFIC PRIORITY SECTION =====
        # Insert at the beginning after the header
        priority_heading = doc.add_paragraph()
        priority_heading.text = f"\n🎯 PRIORITY FOCUS AREA FOR THIS CALL"
        priority_heading.style = 'Heading 1'
        priority_heading.insert_paragraph_before()

        priority_content = doc.add_paragraph()
        priority_content.text = f"""
Research Area: {research_area}

Evidence from National Data (2023-2025):
{evidence_summary}

Number of Internship Positions Available: {internship_count}
Internship Duration: {duration_months} months
Primary Focus: Data analysis, monitoring, and evidence-based intervention design for {research_area}

This internship call specifically targets students interested in addressing {research_area} through data-driven approaches. Selected interns will work directly with NIH researchers and hospital data to develop actionable insights.
"""

        # Move this section to position 2 (after title, before Background)
        # We'll insert it by adding a page break after it

        # ===== SAVE DOCUMENT =====
        safe_filename = research_area.replace(" ", "_").replace("/", "-").replace(":", "")
        output_filename = f"NIH_Internship_Call_{safe_filename}_{today.strftime('%Y%m%d')}.docx"
        output_path = os.path.join(WHO_PROPOSALS_DIR, output_filename)

        doc.save(output_path)

        log_trace("internship_document_generated", {
            "research_area": research_area,
            "file_path": output_path,
            "internship_count": internship_count,
            "duration_months": duration_months
        })

        return {
            "status": "success",
            "message": "Internship call document generated and filled successfully",
            "file_path": output_path,
            "filename": output_filename,
            "research_area": research_area,
            "internship_count": internship_count,
            "duration_months": duration_months,
            "application_deadline": deadline.strftime("%B %d, %Y"),
            "start_date": start_date.strftime("%B %d, %Y"),
            "generated_at": today.isoformat()
        }

    except Exception as e:
        log_trace("internship_document_error", {"error": str(e)})
        return {
            "status": "error",
            "message": f"Failed to generate internship document: {str(e)}"
        }

# =============================================
# RESEARCH PRIORITY ANALYSIS
# =============================================

@mcp.tool()
@traced_tool
def identify_research_priorities(national_data_all_depts: dict) -> dict:
    """
    Analyze national data across all departments to identify research priorities.
    Flags high/medium/low priority areas based on evidence.

    Args:
        national_data_all_depts: Dict with all department national statistics

    Returns:
        dict with prioritized research areas and recommendations
    """

    priorities = {
        "high_priority": [],
        "medium_priority": [],
        "low_priority": [],
        "recommendations": {
            "urgent_interventions": [],
            "research_needed": [],
            "who_proposal_topics": []
        }
    }

    # Analyze each department
    for dept, data in national_data_all_depts.get('department_breakdown', {}).items():
        total_cases = data.get('total_cases_national', 0)
        hospitals_reporting = data.get('hospitals_reporting', 0)

        # Priority classification
        if total_cases > 5000 or hospitals_reporting >= 8:
            priority_level = "high_priority"
            priorities["recommendations"]["who_proposal_topics"].append(dept)
        elif total_cases > 2000 or hospitals_reporting >= 5:
            priority_level = "medium_priority"
        else:
            priority_level = "low_priority"

        priorities[priority_level].append({
            "department": dept,
            "total_cases": total_cases,
            "hospitals_reporting": hospitals_reporting,
            "justification": f"{total_cases} cases across {hospitals_reporting} hospitals"
        })

    # Generate recommendations
    if len(priorities["high_priority"]) > 0:
        priorities["recommendations"]["urgent_interventions"] = [
            f"Immediate intervention needed in {priorities['high_priority'][0]['department']}",
            "Scale up screening programs",
            "Increase healthcare worker training"
        ]

        priorities["recommendations"]["research_needed"] = [
            f"Effectiveness studies for {priorities['high_priority'][0]['department']}",
            "Community-based intervention models",
            "Cost-effectiveness analysis"
        ]

    log_trace("research_priorities_identified", {
        "high_priority_count": len(priorities["high_priority"]),
        "medium_priority_count": len(priorities["medium_priority"]),
        "who_proposals_recommended": len(priorities["recommendations"]["who_proposal_topics"])
    })

    return priorities


# =============================================
# RUN SERVER
# =============================================

if __name__ == "__main__":
    debug_log("🚀 Starting R&D MCP Tools Server...")
    mcp.run()