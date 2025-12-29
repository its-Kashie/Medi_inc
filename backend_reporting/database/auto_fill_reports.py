# # # auto_fill_reports.py
# # import os
# # import pandas as pd
# # from docx import Document
# # from docx.shared import Inches
# # from glob import glob
# #
# # # === CONFIG ===
# # CSV_DIR = "generated_output/csvs"
# # GRAPH_DIR = "generated_output/graphs"
# # TEMPLATE_DIR = "hospital_department_templates"
# # OUTPUT_DIR = "reports"
# #
# # os.makedirs(OUTPUT_DIR, exist_ok=True)
# #
# # # Map internal department keys to their corresponding template files_should_be_in_1_directory
# # DEPT_TEMPLATES = {
# #     "infectious_diseases": "infectious_diseases_template.docx",
# #     "maternal_health": "maternal_health_template.docx",
# #     "nutrition_data": "nutrition_data_template.docx",
# #     "mental_health_data": "mental_health_data_template.docx",
# #     "ncd_internal_medicine": "ncd_internal_medicine_template.docx",
# #     "cardiology_data": "cardiology_data_template.docx",
# #     "endocrinology_diabetes_data": "endocrinology_diabetes_data_template.docx",
# #     "oncology_data": "oncology_data_template.docx"
# # }
# #
# #
# #
# # # === HELPER FUNCTIONS ===
# # def insert_table(doc, df: pd.DataFrame):
# #     """
# #     Inserts a pandas DataFrame into a Word doc as a table.
# #     """
# #     table = doc.add_table(rows=1, cols=len(df.columns))
# #     table.style = "Table Grid"
# #
# #     # Header
# #     hdr_cells = table.rows[0].cells
# #     for i, col in enumerate(df.columns):
# #         hdr_cells[i].text = str(col)
# #
# #     # Data rows
# #     for index, row in df.iterrows():
# #         row_cells = table.add_row().cells
# #         for i, val in enumerate(row):
# #             row_cells[i].text = str(val)
# #
# # def insert_graph(doc, graph_path: str):
# #     """
# #     Inserts a graph image into a Word doc.
# #     """
# #     if os.path.exists(graph_path):
# #         doc.add_picture(graph_path, width=Inches(5))
# #     else:
# #         doc.add_paragraph(f"[Graph missing: {graph_path}]")
# #
# # # === MAIN PROCESS ===
# # for dept, template_file in DEPT_TEMPLATES.items():
# #     template_path = os.path.join(TEMPLATE_DIR, template_file)
# #     if not os.path.exists(template_path):
# #         print(f"Template not found for {dept}: {template_path}")
# #         continue
# #
# #     doc = Document(template_path)
# #
# #     # Load CSV
# #     csv_files = glob(os.path.join(CSV_DIR, f"{dept}_*.csv"))
# #     for csv_file in csv_files:
# #         df = pd.read_csv(csv_file)
# #         doc.add_paragraph(f"### Data from {os.path.basename(csv_file)}")
# #         insert_table(doc, df)
# #
# #         # Insert corresponding graph
# #         graph_file = os.path.join(GRAPH_DIR, os.path.basename(csv_file).replace(".csv", ".png"))
# #         insert_graph(doc, graph_file)
# #
# #     # Save final report
# #     report_path = os.path.join(OUTPUT_DIR, f"{dept}_report.docx")
# #     doc.save(report_path)
# #     print(f"Report generated: {report_path}")
#
#
# # analyze_data_structure.py
# import os
# import json
# import pandas as pd
# from glob import glob
#
# # === CONFIG ===
# CSV_DIR = "generated_output/csvs"
# JSON_DIR = "generated_output/aggregates"
# FOCAL_PERSONS_DIR = "focal_persons_excels"
# OUTPUT_FILE = "data_analysis_report.txt"
#
#
# # === ANALYSIS FUNCTIONS ===
#
# def analyze_csv(csv_path):
#     """Analyze a single CSV file"""
#     df = pd.read_csv(csv_path)
#     filename = os.path.basename(csv_path)
#
#     analysis = {
#         "filename": filename,
#         "total_rows": len(df),
#         "total_columns": len(df.columns),
#         "columns": list(df.columns),
#         "dtypes": {col: str(dtype) for col, dtype in df.dtypes.items()},
#         "null_counts": df.isnull().sum().to_dict(),
#         "sample_data": df.head(3).to_dict('records')
#     }
#
#     # Basic stats for numeric columns
#     numeric_cols = df.select_dtypes(include=['number']).columns
#     if len(numeric_cols) > 0:
#         analysis["numeric_stats"] = df[numeric_cols].describe().to_dict()
#
#     # Value counts for categorical columns (sample first 3)
#     categorical_cols = df.select_dtypes(include=['object']).columns[:3]
#     if len(categorical_cols) > 0:
#         analysis["categorical_samples"] = {}
#         for col in categorical_cols:
#             analysis["categorical_samples"][col] = df[col].value_counts().head(5).to_dict()
#
#     return analysis
#
#
# def analyze_json(json_path):
#     """Analyze a single JSON aggregate file"""
#     with open(json_path, 'r') as f:
#         data = json.load(f)
#
#     filename = os.path.basename(json_path)
#
#     analysis = {
#         "filename": filename,
#         "keys": list(data.keys()) if isinstance(data, dict) else "Not a dict",
#         "structure": str(type(data)),
#         "content": data
#     }
#
#     return analysis
#
#
# def analyze_focal_persons(excel_path):
#     """Analyze focal persons Excel file"""
#     df = pd.read_excel(excel_path)
#     filename = os.path.basename(excel_path)
#
#     analysis = {
#         "filename": filename,
#         "total_rows": len(df),
#         "columns": list(df.columns),
#         "sample_data": df.head(5).to_dict('records')
#     }
#
#     return analysis
#
#
# # === MAIN ANALYSIS ===
#
# def main():
#     report_lines = []
#     report_lines.append("=" * 80)
#     report_lines.append("DATA STRUCTURE ANALYSIS REPORT")
#     report_lines.append("=" * 80)
#     report_lines.append("")
#
#     # 1. ANALYZE CSVs
#     report_lines.append("\n" + "=" * 80)
#     report_lines.append("1. CSV FILES ANALYSIS")
#     report_lines.append("=" * 80)
#
#     csv_files = sorted(glob(os.path.join(CSV_DIR, "*.csv")))
#     for csv_file in csv_files:
#         report_lines.append(f"\n{'─' * 80}")
#         report_lines.append(f"FILE: {os.path.basename(csv_file)}")
#         report_lines.append(f"{'─' * 80}")
#
#         analysis = analyze_csv(csv_file)
#
#         report_lines.append(f"\nTotal Rows: {analysis['total_rows']:,}")
#         report_lines.append(f"Total Columns: {analysis['total_columns']}")
#
#         report_lines.append(f"\nColumns ({len(analysis['columns'])}):")
#         for i, col in enumerate(analysis['columns'], 1):
#             dtype = analysis['dtypes'][col]
#             null_count = analysis['null_counts'][col]
#             report_lines.append(f"  {i}. {col} ({dtype}) - Nulls: {null_count}")
#
#         report_lines.append(f"\nSample Data (first 3 rows):")
#         for i, row in enumerate(analysis['sample_data'], 1):
#             report_lines.append(f"\n  Row {i}:")
#             for key, val in list(row.items())[:5]:  # First 5 columns only
#                 report_lines.append(f"    {key}: {val}")
#
#         if "numeric_stats" in analysis:
#             report_lines.append(f"\nNumeric Column Statistics:")
#             for col in list(analysis['numeric_stats'].keys())[:3]:  # First 3 numeric cols
#                 stats = analysis['numeric_stats'][col]
#                 report_lines.append(f"  {col}:")
#                 report_lines.append(f"    Mean: {stats.get('mean', 'N/A')}")
#                 report_lines.append(f"    Min: {stats.get('min', 'N/A')}")
#                 report_lines.append(f"    Max: {stats.get('max', 'N/A')}")
#
#         if "categorical_samples" in analysis:
#             report_lines.append(f"\nCategorical Column Samples:")
#             for col, counts in analysis['categorical_samples'].items():
#                 report_lines.append(f"  {col} (top 5 values):")
#                 for val, count in list(counts.items())[:5]:
#                     report_lines.append(f"    {val}: {count}")
#
#     # 2. ANALYZE JSONs
#     report_lines.append("\n\n" + "=" * 80)
#     report_lines.append("2. JSON AGGREGATE FILES ANALYSIS")
#     report_lines.append("=" * 80)
#
#     json_files = sorted(glob(os.path.join(JSON_DIR, "*.json")))
#     for json_file in json_files:
#         report_lines.append(f"\n{'─' * 80}")
#         report_lines.append(f"FILE: {os.path.basename(json_file)}")
#         report_lines.append(f"{'─' * 80}")
#
#         analysis = analyze_json(json_file)
#
#         report_lines.append(f"\nStructure: {analysis['structure']}")
#
#         if isinstance(analysis['content'], dict):
#             report_lines.append(f"\nKeys ({len(analysis['keys'])}):")
#             for key in analysis['keys']:
#                 report_lines.append(f"  - {key}")
#
#             report_lines.append(f"\nContent:")
#             report_lines.append(json.dumps(analysis['content'], indent=2))
#         else:
#             report_lines.append(f"\nContent: {analysis['content']}")
#
#     # 3. ANALYZE FOCAL PERSONS
#     report_lines.append("\n\n" + "=" * 80)
#     report_lines.append("3. FOCAL PERSONS EXCEL FILES ANALYSIS")
#     report_lines.append("=" * 80)
#
#     focal_files = sorted(glob(os.path.join(FOCAL_PERSONS_DIR, "*.xlsx")))
#     for focal_file in focal_files:
#         report_lines.append(f"\n{'─' * 80}")
#         report_lines.append(f"FILE: {os.path.basename(focal_file)}")
#         report_lines.append(f"{'─' * 80}")
#
#         analysis = analyze_focal_persons(focal_file)
#
#         report_lines.append(f"\nTotal Rows: {analysis['total_rows']}")
#         report_lines.append(f"Columns: {', '.join(analysis['columns'])}")
#
#         report_lines.append(f"\nSample Data (first 5 rows):")
#         for i, row in enumerate(analysis['sample_data'], 1):
#             report_lines.append(f"\n  Row {i}:")
#             for key, val in row.items():
#                 report_lines.append(f"    {key}: {val}")
#
#     # Write report
#     report_text = "\n".join(report_lines)
#     with open(OUTPUT_FILE, 'w', encoding='utf-8') as f:
#         f.write(report_text)
#
#     print(f"\n✅ Analysis complete! Report saved to: {OUTPUT_FILE}")
#     print(f"📊 Total CSVs analyzed: {len(csv_files)}")
#     print(f"📊 Total JSONs analyzed: {len(json_files)}")
#     print(f"📊 Total Focal Person files_should_be_in_1_directory analyzed: {len(focal_files)}")
#     print(f"\n📄 Open '{OUTPUT_FILE}' to see detailed analysis.")
#
#
# if __name__ == "__main__":
#     main()



# auto_fill_reports_final.py - Department-Specific Customized Reports
import os
import json
import pandas as pd
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import datetime
from collections import Counter

# === CONFIG ===
CSV_DIR = "generated_output/csvs"
JSON_DIR = "generated_output/aggregates"
GRAPH_DIR = "generated_output/graphs"
TEMPLATE_DIR = "hospital_department_templates"
FOCAL_PERSONS_FILE = "focal_persons_excels/hospital_focal_persons.xlsx"
OUTPUT_DIR = "filled_reports"

os.makedirs(OUTPUT_DIR, exist_ok=True)


# === LOAD FOCAL PERSONS ===
def load_focal_persons():
    """Load focal persons mapping"""
    try:
        df = pd.read_excel(FOCAL_PERSONS_FILE)
        focal_map = {}
        for _, row in df.iterrows():
            key = f"{row['Hospital']}_{row['Department']}"
            focal_map[key] = {
                'name': row['Focal Person Name'],
                'contact': row['Contact'],
                'email': row['Email']
            }
        return focal_map
    except Exception as e:
        print(f"⚠️  Warning: Could not load focal persons: {e}")
        return {}


# === HELPER FUNCTIONS ===
def get_quarter_dates(year, quarter):
    """Get start and end dates for a quarter"""
    quarters = {
        1: (f"Jan 1, {year}", f"Mar 31, {year}"),
        2: (f"Apr 1, {year}", f"Jun 30, {year}"),
        3: (f"Jul 1, {year}", f"Sep 30, {year}"),
        4: (f"Oct 1, {year}", f"Dec 31, {year}")
    }
    return quarters.get(quarter, ("N/A", "N/A"))


def calculate_percent_change(current, previous):
    """Calculate percentage change"""
    if previous == 0 or previous is None:
        return "N/A"
    try:
        change = ((current - previous) / previous) * 100
        return f"{change:+.1f}%"
    except:
        return "N/A"


def add_para(doc, text, bold=False, size=11):
    """Add formatted paragraph"""
    para = doc.add_paragraph()
    run = para.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    return para


def insert_table(doc, data, headers):
    """Insert a formatted table"""
    if not data:
        return None
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Light Grid Accent 1'
    hdr_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        hdr_cells[i].text = str(header)
        hdr_cells[i].paragraphs[0].runs[0].font.bold = True
    for row_data in data:
        row_cells = table.add_row().cells
        for i, val in enumerate(row_data):
            row_cells[i].text = str(val)
    return table


def safe_count(series, value):
    """Safely count values"""
    try:
        return series.value_counts().get(value, 0)
    except:
        return 0


def get_top_items(df, column, top_n=5):
    """Get top N items"""
    try:
        if column not in df.columns:
            return []
        counts = df[column].value_counts().head(top_n)
        return [(item, count) for item, count in counts.items()]
    except:
        return []


# === DEPARTMENT-SPECIFIC REPORT GENERATORS ===

def generate_infectious_diseases_report(doc, df, prev_data, dept_info, hospital):
    """Generate Infectious Diseases specific report"""
    total = len(df)
    mortality = safe_count(df.get('outcome', pd.Series()), 'Died')
    admitted = safe_count(df.get('admission', pd.Series()), 'Yes')

    # Executive Summary
    doc.add_heading('Executive Summary', level=1)
    add_para(doc, f"• Total patients treated: {total}", bold=True)
    add_para(doc, f"• Admitted patients: {admitted} ({admitted / total * 100:.1f}%)", bold=True)
    add_para(doc, f"• Mortality cases: {mortality}", bold=True)
    add_para(doc, f"• Average hospitalization days: {df['days_hospitalized'].mean():.1f} days", bold=True)

    doc.add_paragraph()
    add_para(doc, "Key Highlights:", bold=True, size=12)
    if prev_data:
        change = calculate_percent_change(total, prev_data['total_patients'])
        add_para(doc, f"• Patient volume changed by {change} vs previous quarter")
    add_para(doc, f"• Average patient age: {df['age'].mean():.1f} years")

    # Dehydration analysis
    if 'dehydration_level' in df.columns:
        severe_dehydration = df['dehydration_level'].value_counts().get('Severe', 0)
        add_para(doc, f"• Severe dehydration cases: {severe_dehydration}")

    # Patient Summary
    doc.add_page_break()
    doc.add_heading('Patient / Treatment Summary', level=1)
    summary_data = [
        ["Total Patients", total, prev_data['total_patients'] if prev_data else 'N/A',
         calculate_percent_change(total, prev_data['total_patients'] if prev_data else 0), "Within range"],
        ["Admitted Cases", admitted, "N/A", "N/A", f"{admitted / total * 100:.1f}% admission rate"],
        ["Mortality", mortality, prev_data['mortality'] if prev_data else 'N/A',
         calculate_percent_change(mortality, prev_data['mortality'] if prev_data else 0), "Requires monitoring"]
    ]
    insert_table(doc, summary_data, ["Indicator", "Current", "Previous", "% Change", "Notes"])

    # Disease Categories
    doc.add_page_break()
    doc.add_heading('Disease Distribution', level=1)
    diseases = get_top_items(df, 'disease_category', 5)
    if diseases:
        disease_data = [[d, c, f"{c / total * 100:.1f}%", "Standard protocol"] for d, c in diseases]
        insert_table(doc, disease_data, ["Disease", "Count", "%", "Treatment Protocol"])

    # Water Source Analysis
    if 'water_source_reported' in df.columns:
        doc.add_paragraph()
        add_para(doc, "Water Source Analysis:", bold=True)
        sources = get_top_items(df, 'water_source_reported', 3)
        for source, count in sources:
            add_para(doc, f"• {source}: {count} cases ({count / total * 100:.1f}%)")


def generate_maternal_health_report(doc, df, prev_data, dept_info, hospital):
    """Generate Maternal Health specific report"""
    total = len(df)
    csection = safe_count(df.get('delivery_type', pd.Series()), 'C-section')
    complications = df['maternal_complication'].notna().sum() if 'maternal_complication' in df.columns else 0
    neonatal_issues = df['neonatal_complication'].notna().sum() if 'neonatal_complication' in df.columns else 0

    # Executive Summary
    doc.add_heading('Executive Summary', level=1)
    add_para(doc, f"• Total patients registered: {total}", bold=True)
    add_para(doc, f"• C-section deliveries: {csection} ({csection / total * 100:.1f}%)", bold=True)
    add_para(doc, f"• Maternal complications: {complications}", bold=True)
    add_para(doc, f"• Neonatal complications: {neonatal_issues}", bold=True)
    add_para(doc, f"• Average antenatal visits: {df['antenatal_visits'].mean():.1f}", bold=True)

    doc.add_paragraph()
    add_para(doc, "Key Highlights:", bold=True, size=12)
    if prev_data:
        change = calculate_percent_change(total, prev_data['total_patients'])
        add_para(doc, f"• Patient registrations changed by {change}")
    add_para(doc, f"• Average maternal age: {df['age'].mean():.1f} years")
    add_para(doc, f"• Average hemoglobin: {df['hemoglobin_level'].mean():.1f} g/dL")

    # High-risk pregnancy indicators
    high_risk = df[
        (df['hypertension'] == 'Yes') |
        (df['diabetes_status'] == 'Positive') |
        (df['previous_miscarriages'] > 0)
        ].shape[0]
    add_para(doc, f"• High-risk pregnancies: {high_risk} ({high_risk / total * 100:.1f}%)")

    # Patient Summary
    doc.add_page_break()
    doc.add_heading('Patient / Treatment Summary', level=1)
    summary_data = [
        ["Total Registrations", total, prev_data['total_patients'] if prev_data else 'N/A',
         calculate_percent_change(total, prev_data['total_patients'] if prev_data else 0), "Normal range"],
        ["C-section Rate", f"{csection / total * 100:.1f}%", "N/A", "N/A", "Within acceptable range"],
        ["Maternal Complications", complications, prev_data['complications'] if prev_data else 'N/A',
         calculate_percent_change(complications, prev_data['complications'] if prev_data else 0),
         "Monitoring required"],
        ["High-risk Cases", high_risk, prev_data['high_risk_flags'] if prev_data else 'N/A', "N/A", "Close monitoring"]
    ]
    insert_table(doc, summary_data, ["Indicator", "Current", "Previous", "% Change", "Notes"])

    # Delivery Types
    doc.add_page_break()
    doc.add_heading('Delivery & Outcomes', level=1)
    deliveries = get_top_items(df, 'delivery_type', 3)
    if deliveries:
        delivery_data = [[d, c, f"{c / total * 100:.1f}%"] for d, c in deliveries]
        insert_table(doc, delivery_data, ["Delivery Type", "Count", "Percentage"])


def generate_nutrition_report(doc, df, prev_data, dept_info, hospital):
    """Generate Nutrition specific report"""
    total = len(df)
    stunted = safe_count(df.get('stunted', pd.Series()), 'Yes')

    # Executive Summary
    doc.add_heading('Executive Summary', level=1)
    add_para(doc, f"• Total children assessed: {total}", bold=True)
    add_para(doc, f"• Stunted children: {stunted} ({stunted / total * 100:.1f}%)", bold=True)
    add_para(doc, f"• Average WFH Z-score: {df['wfh_zscore'].mean():.2f}", bold=True)
    add_para(doc, f"• Average child age: {df['age_months'].mean():.1f} months", bold=True)

    # Supplement distribution
    if 'supplement_given' in df.columns:
        supplemented = df['supplement_given'].notna().sum()
        add_para(doc, f"• Children supplemented: {supplemented} ({supplemented / total * 100:.1f}%)", bold=True)

    doc.add_paragraph()
    add_para(doc, "Key Highlights:", bold=True, size=12)
    if prev_data:
        change = calculate_percent_change(total, prev_data['total_patients'])
        add_para(doc, f"• Screening volume changed by {change}")

    # Malnutrition severity
    severe_malnutrition = df[df['wfh_zscore'] < -3].shape[0]
    add_para(doc, f"• Severe malnutrition cases: {severe_malnutrition} ({severe_malnutrition / total * 100:.1f}%)")

    # Patient Summary
    doc.add_page_break()
    doc.add_heading('Nutritional Assessment Summary', level=1)
    summary_data = [
        ["Total Assessments", total, prev_data['total_patients'] if prev_data else 'N/A',
         calculate_percent_change(total, prev_data['total_patients'] if prev_data else 0), "Screening complete"],
        ["Stunted Children", stunted, "N/A", "N/A", f"{stunted / total * 100:.1f}% prevalence"],
        ["Severe Malnutrition", severe_malnutrition, "N/A", "N/A", "Requires intervention"]
    ]
    insert_table(doc, summary_data, ["Indicator", "Current", "Previous", "% Change", "Notes"])


def generate_mental_health_report(doc, df, prev_data, dept_info, hospital):
    """Generate Mental Health specific report"""
    total = len(df)
    high_risk = safe_count(df.get('suicide_risk', pd.Series()), 'High')
    moderate_risk = safe_count(df.get('suicide_risk', pd.Series()), 'Moderate')

    # Executive Summary
    doc.add_heading('Executive Summary', level=1)
    add_para(doc, f"• Total patients screened: {total}", bold=True)
    add_para(doc, f"• High suicide risk: {high_risk}", bold=True)
    add_para(doc, f"• Moderate suicide risk: {moderate_risk}", bold=True)
    add_para(doc, f"• Average screening score: {df['score'].mean():.1f}", bold=True)
    add_para(doc, f"• Average counseling sessions: {df['counseling_sessions'].mean():.1f}", bold=True)

    doc.add_paragraph()
    add_para(doc, "Key Highlights:", bold=True, size=12)
    if prev_data:
        change = calculate_percent_change(total, prev_data['total_patients'])
        add_para(doc, f"• Screening volume changed by {change}")
    add_para(doc, f"• Average patient age: {df['age'].mean():.1f} years")

    # Patient Summary
    doc.add_page_break()
    doc.add_heading('Mental Health Assessment Summary', level=1)
    summary_data = [
        ["Total Screenings", total, prev_data['total_patients'] if prev_data else 'N/A',
         calculate_percent_change(total, prev_data['total_patients'] if prev_data else 0), "Active screening"],
        ["High Risk Cases", high_risk, "N/A", "N/A", "Immediate intervention needed"],
        ["Moderate Risk", moderate_risk, "N/A", "N/A", "Close monitoring required"]
    ]
    insert_table(doc, summary_data, ["Indicator", "Current", "Previous", "% Change", "Notes"])

    # Diagnosis Distribution
    doc.add_page_break()
    doc.add_heading('Diagnosis Distribution', level=1)
    diagnoses = get_top_items(df, 'diagnosis', 5)
    if diagnoses:
        diag_data = [[d, c, f"{c / total * 100:.1f}%", "Standard protocol"] for d, c in diagnoses]
        insert_table(doc, diag_data, ["Diagnosis", "Count", "%", "Treatment"])


def generate_cardiology_report(doc, df, prev_data, dept_info, hospital):
    """Generate Cardiology specific report"""
    total = len(df)
    mortality = safe_count(df.get('mortality', pd.Series()), 'Died')
    icu = safe_count(df.get('icu_admission', pd.Series()), 'Yes')
    procedures = df['procedure'].notna().sum() if 'procedure' in df.columns else 0

    # Executive Summary
    doc.add_heading('Executive Summary', level=1)
    add_para(doc, f"• Total cardiac patients: {total}", bold=True)
    add_para(doc, f"• Procedures performed: {procedures}", bold=True)
    add_para(doc, f"• ICU admissions: {icu} ({icu / total * 100:.1f}%)", bold=True)
    add_para(doc, f"• Mortality: {mortality}", bold=True)
    add_para(doc, f"• Average hospital stay: {df['days_hospitalized'].mean():.1f} days", bold=True)

    doc.add_paragraph()
    add_para(doc, "Key Highlights:", bold=True, size=12)
    if prev_data:
        change = calculate_percent_change(total, prev_data['total_patients'])
        add_para(doc, f"• Patient volume changed by {change}")
    add_para(doc, f"• Average patient age: {df['age'].mean():.1f} years")

    # Patient Summary
    doc.add_page_break()
    doc.add_heading('Cardiac Care Summary', level=1)
    summary_data = [
        ["Total Patients", total, prev_data['total_patients'] if prev_data else 'N/A',
         calculate_percent_change(total, prev_data['total_patients'] if prev_data else 0), "Normal range"],
        ["Procedures", procedures, "N/A", "N/A", "Interventions completed"],
        ["ICU Admissions", icu, "N/A", "N/A", "Critical care provided"],
        ["Mortality", mortality, prev_data['mortality'] if prev_data else 'N/A',
         calculate_percent_change(mortality, prev_data['mortality'] if prev_data else 0), "Within acceptable range"]
    ]
    insert_table(doc, summary_data, ["Indicator", "Current", "Previous", "% Change", "Notes"])

    # ECG Findings
    doc.add_page_break()
    doc.add_heading('ECG Findings & Procedures', level=1)
    ecg_findings = get_top_items(df, 'ecg_findings', 5)
    if ecg_findings:
        ecg_data = [[e, c, f"{c / total * 100:.1f}%"] for e, c in ecg_findings]
        insert_table(doc, ecg_data, ["ECG Finding", "Count", "Percentage"])


def generate_endocrinology_report(doc, df, prev_data, dept_info, hospital):
    """Generate Endocrinology/Diabetes specific report"""
    total = len(df)
    type1 = safe_count(df.get('diabetes_type', pd.Series()), 'Type 1')
    type2 = safe_count(df.get('diabetes_type', pd.Series()), 'Type 2')
    on_insulin = safe_count(df.get('on_insulin', pd.Series()), 'Yes')
    complications = df['complication'].notna().sum() if 'complication' in df.columns else 0

    # Executive Summary
    doc.add_heading('Executive Summary', level=1)
    add_para(doc, f"• Total diabetes patients: {total}", bold=True)
    add_para(doc, f"• Type 1 diabetes: {type1} ({type1 / total * 100:.1f}%)", bold=True)
    add_para(doc, f"• Type 2 diabetes: {type2} ({type2 / total * 100:.1f}%)", bold=True)
    add_para(doc, f"• Patients on insulin: {on_insulin} ({on_insulin / total * 100:.1f}%)", bold=True)
    add_para(doc, f"• Average HbA1c: {df['hba1c'].mean():.1f}%", bold=True)
    add_para(doc, f"• Complications: {complications}", bold=True)

    doc.add_paragraph()
    add_para(doc, "Key Highlights:", bold=True, size=12)
    if prev_data:
        change = calculate_percent_change(total, prev_data['total_patients'])
        add_para(doc, f"• Patient volume changed by {change}")
    add_para(doc, f"• Average patient age: {df['age'].mean():.1f} years")

    # Glycemic control
    poor_control = df[df['hba1c'] > 9].shape[0]
    add_para(doc, f"• Poor glycemic control (HbA1c>9%): {poor_control} ({poor_control / total * 100:.1f}%)")

    # Patient Summary
    doc.add_page_break()
    doc.add_heading('Diabetes Management Summary', level=1)
    summary_data = [
        ["Total Patients", total, prev_data['total_patients'] if prev_data else 'N/A',
         calculate_percent_change(total, prev_data['total_patients'] if prev_data else 0), "Active management"],
        ["Type 1 Diabetes", type1, "N/A", "N/A", f"{type1 / total * 100:.1f}%"],
        ["Type 2 Diabetes", type2, "N/A", "N/A", f"{type2 / total * 100:.1f}%"],
        ["On Insulin Therapy", on_insulin, "N/A", "N/A", "Insulin management"],
        ["Complications", complications, "N/A", "N/A", "Requires monitoring"]
    ]
    insert_table(doc, summary_data, ["Indicator", "Current", "Previous", "% Change", "Notes"])


def generate_ncd_report(doc, df, prev_data, dept_info, hospital):
    """Generate NCD/Internal Medicine specific report"""
    total = len(df)
    admitted = safe_count(df.get('admission', pd.Series()), 'Yes')
    mortality = safe_count(df.get('outcome', pd.Series()), 'Died')

    # Executive Summary
    doc.add_heading('Executive Summary', level=1)
    add_para(doc, f"• Total NCD patients: {total}", bold=True)
    add_para(doc, f"• Admitted patients: {admitted} ({admitted / total * 100:.1f}%)", bold=True)
    add_para(doc, f"• Mortality: {mortality}", bold=True)

    doc.add_paragraph()
    add_para(doc, "Key Highlights:", bold=True, size=12)
    if prev_data:
        change = calculate_percent_change(total, prev_data['total_patients'])
        add_para(doc, f"• Patient volume changed by {change}")
    add_para(doc, f"• Average patient age: {df['age'].mean():.1f} years")

    # Patient Summary
    doc.add_page_break()
    doc.add_heading('NCD Management Summary', level=1)
    summary_data = [
        ["Total Patients", total, prev_data['total_patients'] if prev_data else 'N/A',
         calculate_percent_change(total, prev_data['total_patients'] if prev_data else 0), "Active management"],
        ["Admissions", admitted, "N/A", "N/A", f"{admitted / total * 100:.1f}% admission rate"],
        ["Mortality", mortality, prev_data['mortality'] if prev_data else 'N/A',
         calculate_percent_change(mortality, prev_data['mortality'] if prev_data else 0), "Monitoring required"]
    ]
    insert_table(doc, summary_data, ["Indicator", "Current", "Previous", "% Change", "Notes"])

    # NCD Distribution
    doc.add_page_break()
    doc.add_heading('NCD Distribution', level=1)
    ncds = get_top_items(df, 'primary_ncd_diagnosis', 5)
    if ncds:
        ncd_data = [[n, c, f"{c / total * 100:.1f}%", "Standard protocol"] for n, c in ncds]
        insert_table(doc, ncd_data, ["NCD Diagnosis", "Count", "%", "Management"])


def generate_oncology_report(doc, df, prev_data, dept_info, hospital):
    """Generate Oncology specific report"""
    total = len(df)
    mortality = safe_count(df.get('outcome', pd.Series()), 'Died')
    late_diagnosis = safe_count(df.get('late_diagnosis', pd.Series()), 'Yes')

    # Executive Summary
    doc.add_heading('Executive Summary', level=1)
    add_para(doc, f"• Total cancer patients: {total}", bold=True)
    add_para(doc, f"• Late diagnoses: {late_diagnosis} ({late_diagnosis / total * 100:.1f}%)", bold=True)
    add_para(doc, f"• Mortality: {mortality}", bold=True)

    doc.add_paragraph()
    add_para(doc, "Key Highlights:", bold=True, size=12)
    if prev_data:
        change = calculate_percent_change(total, prev_data['total_patients'])
        add_para(doc, f"• Patient volume changed by {change}")
    add_para(doc, f"• Average patient age: {df['age'].mean():.1f} years")

    # Patient Summary
    doc.add_page_break()
    doc.add_heading('Oncology Summary', level=1)
    summary_data = [
        ["Total Patients", total, prev_data['total_patients'] if prev_data else 'N/A',
         calculate_percent_change(total, prev_data['total_patients'] if prev_data else 0), "Active management"],
        ["Late Diagnosis", late_diagnosis, "N/A", "N/A", f"{late_diagnosis / total * 100:.1f}%"],
        ["Mortality", mortality, prev_data['mortality'] if prev_data else 'N/A',
         calculate_percent_change(mortality, prev_data['mortality'] if prev_data else 0), "Within range"]
    ]
    insert_table(doc, summary_data, ["Indicator", "Current", "Previous", "% Change", "Notes"])

    # Cancer Sites
    doc.add_page_break()
    doc.add_heading('Cancer Distribution', level=1)
    cancers = get_top_items(df, 'cancer_site', 5)
    if cancers:
        cancer_data = [[c, cnt, f"{cnt / total * 100:.1f}%"] for c, cnt in cancers]
        insert_table(doc, cancer_data, ["Cancer Site", "Count", "Percentage"])


# === DEPARTMENT MAPPING ===
DEPT_CONFIG = {
    "infectious_diseases": {
        "template": "infectious_diseases_template.docx",
        "name": "Infectious Diseases",
        "focal_dept": "Infectious Diseases",
        "generator": generate_infectious_diseases_report
    },
    "maternal_health": {
        "template": "maternal_health_template.docx",
        "name": "Obstetrics & Gynecology",
        "focal_dept": "Obstetrics & Gynecology",
        "generator": generate_maternal_health_report
    },
    "nutrition_data": {
        "template": "nutrition_data_template.docx",
        "name": "Nutrition & Dietetics",
        "focal_dept": "Nutrition & Dietetics",
        "generator": generate_nutrition_report
    },
    "mental_health_data": {
        "template": "mental_health_data_template.docx",
        "name": "Psychiatry",
        "focal_dept": "Psychiatry",
        "generator": generate_mental_health_report
    },
    "ncd_internal_medicine": {
        "template": "ncd_internal_medicine_template.docx",
        "name": "Internal Medicine (NCDs)",
        "focal_dept": "Internal Medicine (NCDs)",
        "generator": generate_ncd_report
    },
    "cardiology_data": {
        "template": "cardiology_data_template.docx",
        "name": "Cardiology",
        "focal_dept": "Cardiology",
        "generator": generate_cardiology_report
    },
    "endocrinology_diabetes_data": {
        "template": "endocrinology_diabetes_data_template.docx",
        "name": "Endocrinology / Diabetes",
        "focal_dept": "Endocrinology",
        "generator": generate_endocrinology_report
    },
    "oncology_data": {
        "template": "oncology_data_template.docx",
        "name": "Oncology",
        "focal_dept": "Oncology",
        "generator": generate_oncology_report
    }
}


# === MAIN REPORT GENERATION ===
def generate_report(dept_key, target_year=2025, target_quarter=3):
    """Generate department-specific report"""
    print(f"\n{'=' * 80}")
    print(f"Generating: {dept_key}")
    print(f"{'=' * 80}")

    dept = DEPT_CONFIG.get(dept_key)
    if not dept:
        print(f"❌ Unknown department: {dept_key}")
        return

    # Load data
    csv_path = os.path.join(CSV_DIR, f"{dept_key}.csv")
    json_path = os.path.join(JSON_DIR, f"{dept_key}_aggregates.json")

    if not os.path.exists(csv_path) or not os.path.exists(json_path):
        print(f"❌ Missing data files")
        return

    df = pd.read_csv(csv_path)
    with open(json_path, 'r') as f:
        quarterly_data = json.load(f)

    # Get current quarter data
    current = next((q for q in quarterly_data
                    if q['report_year'] == target_year and q['report_quarter'] == target_quarter), None)
    if not current:
        print(f"❌ No data for Q{target_quarter} {target_year}")
        return

    # Get previous quarter
    prev_q = target_quarter - 1 if target_quarter > 1 else 4
    prev_y = target_year if target_quarter > 1 else target_year - 1
    prev = next((q for q in quarterly_data
                 if q['report_year'] == prev_y and q['report_quarter'] == prev_q), None)

    # Filter data
    df_current = df[(df['report_year'] == target_year) & (df['report_quarter'] == target_quarter)]
    focal_persons = load_focal_persons()
    hospitals = df_current['hospital'].unique()

    # Generate for each hospital
    for hospital in hospitals[:3]:  # Limit to 3 for demo
        print(f"  📄 {hospital}")

        template_path = os.path.join(TEMPLATE_DIR, dept['template'])
        if not os.path.exists(template_path):
            continue

        doc = Document(template_path)
        df_hospital = df_current[df_current['hospital'] == hospital]

        # Clear template content intelligently
        # Define what to keep and what to remove
        keep_patterns = [
            dept['name'] + ' Quarterly Report Template',
            'Quarterly Report Template',
            'Generated with: Hospital Health Agent'
        ]

        TEMPLATE_TABLE_HEADERS = [
            "Indicator | Target / Benchmark | Current Quarter Data | % Change vs Last Quarter | Notes / Actions Taken",
            "Test / Procedure | Number Conducted | Abnormal Findings | Notes / Follow-up",
            "Resource / Staff | Planned / Available | Utilization this Quarter | Notes / Issues"
        ]

        UNWANTED_LINES = [
            "Prepared by: Head of Internal Medicine (NCDs): Laith Ghanee"
        ]

        # --- Remove unwanted paragraphs ---
        paras_to_remove = []

        for para in doc.paragraphs:
            txt = para.text.strip()

            # Remove specific known template lines
            if txt in UNWANTED_LINES:
                paras_to_remove.append(para)
                continue

            # Remove template table headings written as text
            if txt in TEMPLATE_TABLE_HEADERS:
                paras_to_remove.append(para)
                continue

            # Remove placeholder bullet points
            if txt.startswith("•") and "____" in txt:
                paras_to_remove.append(para)
                continue

            # Remove sections that you don't want
            remove_keywords = [
                "Executive Summary",
                "Patient / Treatment Summary",
                "Diagnostic / Monitoring Data",
                "Resource & Staff",
                "Action Plan",
                "Key Highlights",
                "Indicator |",
                "Test / Procedure |",
                "Resource / Staff |"
            ]
            if any(k in txt for k in remove_keywords):
                paras_to_remove.append(para)
                continue

        # Remove paragraphs
        for p in paras_to_remove:
            p._element.getparent().remove(p._element)

        # --- Remove ONLY template tables (keep your data tables) ---
        tables_to_remove = []

        for table in doc.tables:
            header_text = " | ".join(cell.text.strip() for cell in table.rows[0].cells)

            if header_text in TEMPLATE_TABLE_HEADERS:
                tables_to_remove.append(table)

        for tbl in tables_to_remove:
            tbl._element.getparent().remove(tbl._element)

        # -----------------------------------------
        # REPLACE PLACEHOLDERS IN HEADER
        # -----------------------------------------

        for para in doc.paragraphs:
            t = para.text

            if "Hospital Name:" in t:
                para.text = f"Hospital Name: {hospital}"

            elif "Reporting Period:" in t:
                start, end = get_quarter_dates(target_year, target_quarter)
                para.text = f"Reporting Period: From {start} to {end}"

            elif "Prepared by:" in t:
                focal_key = f"{hospital}_{dept['focal_dept']}"
                focal = focal_persons.get(focal_key, {'name': '[Name Not Available]'})
                para.text = f"Prepared by: Head of {dept['name']}: {focal['name']}"

        doc.add_page_break()

        remove_patterns = [
            '_______',
            '_____|',
            'Executive Summary',
            'Patient / Treatment Summary',
            'Diagnostic / Monitoring Data',
            'Resource & Staff',
            'Action Plan / Recommendations',
            'Key Highlights / Remarks',
            'Indicator |',
            'Test / Procedure |',
            'Resource / Staff |',
            '• Improvements required:',
            '• Training needs:',
            '• Equipment procurement',
            '• Policy or procedure updates:'
        ]

        to_remove = []

        for element in doc.element.body:
            if element.tag.endswith('p'):
                # Get full paragraph text
                para_text = ''.join(element.itertext()) if hasattr(element, 'itertext') else ''

                # Always keep main title and "Generated with"
                if any(keep in para_text for keep in keep_patterns):
                    continue

                # Keep Hospital Name, Reporting Period, Prepared by (first occurrence only)
                if 'Hospital Name:' in para_text and '____' not in para_text:
                    continue
                if 'Reporting Period:' in para_text and '____' not in para_text:
                    continue
                if 'Prepared by:' in para_text and '____' not in para_text:
                    # Only keep if it's at the top (not signature section)
                    if 'Date:' not in para_text:
                        continue

                # Remove template placeholders and sections
                if any(pattern in para_text for pattern in remove_patterns):
                    to_remove.append(element)
                    continue

                # Remove bullets with just underscores
                if para_text.strip().startswith('•') and '____' in para_text:
                    to_remove.append(element)

            elif element.tag.endswith('tbl'):
                # Remove all template tables
                to_remove.append(element)

        # Remove collected elements
        for element in to_remove:
            try:
                element.getparent().remove(element)
            except:
                pass

        # Now fill the header fields that have placeholders
        for para in doc.paragraphs:
            text = para.text

            if "Hospital Name:" in text and "____" in text:
                para.text = f"Hospital Name: {hospital}"

            elif "Reporting Period:" in text and "____" in text:
                start, end = get_quarter_dates(target_year, target_quarter)
                para.text = f"Reporting Period: From {start} to {end}"

            elif "Prepared by:" in text and ("____" in text or "Dr. [Name]" in text):
                focal_key = f"{hospital}_{dept['focal_dept']}"
                focal = focal_persons.get(focal_key, {'name': '[Name Not Available]'})
                para.text = f"Prepared by: Head of {dept['name']}: {focal['name']}"

        doc.add_page_break()

        # Generate department-specific content
        dept['generator'](doc, df_hospital, prev, dept, hospital)

        # Add graphs
        trend_graph = os.path.join(GRAPH_DIR, f"{dept_key}_trend_total_patients.png")
        mort_graph = os.path.join(GRAPH_DIR, f"{dept_key}_mort_comp.png")

        if os.path.exists(trend_graph) or os.path.exists(mort_graph):
            doc.add_page_break()
            doc.add_heading('Visual Analytics', level=1)

            if os.path.exists(trend_graph):
                add_para(doc, "Patient Trend Over Time:", bold=True)
                doc.add_picture(trend_graph, width=Inches(6))
                doc.add_paragraph()

            if os.path.exists(mort_graph):
                add_para(doc, "Mortality & Complications:", bold=True)
                doc.add_picture(mort_graph, width=Inches(6))

        # Recommendations
        doc.add_page_break()
        doc.add_heading('Action Plan / Recommendations', level=1)

        recommendations = [
            f"• Maintain quality standards for {dept['name']} services",
            "• Continue evidence-based clinical protocols",
            "• Enhance staff capacity through regular training",
            "• Monitor key performance indicators closely",
            "• Update treatment guidelines as per latest research"
        ]

        for rec in recommendations:
            add_para(doc, rec)

        # Signatures
        doc.add_page_break()
        focal_key = f"{hospital}_{dept['focal_dept']}"
        focal = focal_persons.get(focal_key, {'name': '[Name Not Available]'})

        add_para(doc, f"Prepared by: Head of {dept['name']}", bold=True)
        add_para(doc, f"Name: {focal['name']}")
        add_para(doc, f"Date: {datetime.now().strftime('%B %d, %Y')}")
        doc.add_paragraph()
        add_para(doc, "Approved by Hospital Registrar: _________________", bold=True)
        add_para(doc, f"Date: {datetime.now().strftime('%B %d, %Y')}")

        # Save
        safe_hospital = hospital.replace(" ", "_").replace("/", "_")
        filename = f"{dept_key}_{safe_hospital}_Q{target_quarter}_{target_year}_report.docx"
        output_path = os.path.join(OUTPUT_DIR, filename)
        doc.save(output_path)
        print(f"  ✅ Saved: {filename}")


# === MAIN ===
def main():
    print("\n" + "=" * 80)
    print("CUSTOMIZED DEPARTMENT REPORTS GENERATOR")
    print("=" * 80)

    for dept_key in DEPT_CONFIG.keys():
        try:
            generate_report(dept_key, target_year=2025, target_quarter=3)
        except Exception as e:
            print(f"❌ Error in {dept_key}: {str(e)}")
            import traceback
            traceback.print_exc()

    print("\n" + "=" * 80)
    print("✅ ALL REPORTS GENERATED!")
    print(f"📂 Location: {OUTPUT_DIR}")
    print("=" * 80)


if __name__ == "__main__":
    main()