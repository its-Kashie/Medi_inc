# generate_department_templates.py
from docx import Document
import os

OUTPUT_DIR = "hospital_department_templates"
os.makedirs(OUTPUT_DIR, exist_ok=True)

DEPARTMENTS = {
    "infectious_diseases": "Infectious Diseases (Water-borne)",
    "maternal_health": "Obstetrics & Gynecology (Maternal/Neonatal)",
    "nutrition_data": "Nutrition & Dietetics",
    "mental_health_data": "Psychiatry / Mental Health",
    "ncd_internal_medicine": "Internal Medicine (NCDs)",
    "cardiology_data": "Cardiology",
    "endocrinology_diabetes_data": "Endocrinology / Diabetes",
    "oncology_data": "Oncology"
}

for dept_key, dept_name in DEPARTMENTS.items():
    doc = Document()
    doc.add_heading(f"{dept_name} Quarterly Report Template", 0)
    doc.add_paragraph("Hospital Name: ____________________________")
    doc.add_paragraph("Reporting Period: From ________ to ________")
    doc.add_paragraph(f"Prepared by: Head of {dept_name}: ___________")
    doc.add_paragraph("Generated with: Hospital Health Agent\n")

    doc.add_heading("1. Executive Summary", level=1)
    doc.add_paragraph("• Total patients admitted / treated: _______")
    doc.add_paragraph("• Total procedures / interventions: _______")
    doc.add_paragraph("• ICU admissions / critical cases: _______")
    doc.add_paragraph("• Mortality / complications: _______")
    doc.add_paragraph("Key Highlights / Remarks:\n")

    doc.add_heading("2. Patient / Treatment Summary", level=1)
    doc.add_paragraph("Indicator | Target / Benchmark | Current Quarter Data | % Change vs Last Quarter | Notes / Actions Taken\n")
    doc.add_paragraph("_____|_____|_____|_____|_____\n" * 6)

    doc.add_heading("3. Diagnostic / Monitoring Data", level=1)
    doc.add_paragraph("Test / Procedure | Number Conducted | Abnormal Findings | Notes / Follow-up\n")
    doc.add_paragraph("_____|_____|_____|_____\n" * 5)

    doc.add_heading("4. Resource & Staff Utilization", level=1)
    doc.add_paragraph("Resource / Staff | Planned / Available | Utilization this Quarter | Notes / Issues\n")
    doc.add_paragraph("_____|_____|_____|_____\n" * 5)

    doc.add_heading("5. Action Plan / Recommendations", level=1)
    doc.add_paragraph("• Improvements required: ____________________________________________")
    doc.add_paragraph("• Training needs: _________________________________________________")
    doc.add_paragraph("• Equipment procurement or maintenance: ___________________________")
    doc.add_paragraph("• Policy or procedure updates: _____________________________________\n")

    doc.add_paragraph(f"Prepared by: Head of {dept_name}: ___________________ Date: _______")
    doc.add_paragraph("Approved by Hospital Registrar: _________________ Date: _______")

    # save doc
    doc_path = os.path.join(OUTPUT_DIR, f"{dept_key}_template.docx")
    doc.save(doc_path)
    print(f"Generated template: {doc_path}")
